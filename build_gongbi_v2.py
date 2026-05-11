#!/usr/bin/env python3
"""
Build Tintinalli Ch 56-61 共筆 v2 — using Heading 1-6 styles only,
matching the GI_03.docx reference format exactly.

The template's Heading 1-6 styles are pre-linked to numId=1 -> abstractNumId=1
which auto-generates 一、(一)、1.、(1)、a.、(a) markers via Word's
multilevel list mechanism.

For each Topic we add a fresh numId restart so numbering starts at 一 again.
"""

import zipfile
import shutil
import re
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy import deepcopy

SRC = '/root/.claude/uploads/565389eb-0226-4708-bcd1-3c108a1d5113/5f4dc64c-____.dotx'
TMP = '/tmp/tpl_v2.docx'
OUT = '/home/user/my-first-claude-project-/Tintinalli_Ch56-61_共筆_v2.docx'

# Step 1: dotx -> docx
shutil.copy(SRC, TMP)
with zipfile.ZipFile(TMP, 'r') as zin:
    files = {n: zin.read(n) for n in zin.namelist()}
ct = files['[Content_Types].xml'].decode('utf-8').replace(
    'wordprocessingml.template.main+xml',
    'wordprocessingml.document.main+xml',
)
files['[Content_Types].xml'] = ct.encode('utf-8')

# Step 2: add 6 numIds (100-105) referencing abstractNumId=1 with full level reset
num_xml = files['word/numbering.xml'].decode('utf-8')
extra_nums = ''
for nid in range(100, 106):  # 100..105 -> one per Topic
    overrides = ''.join(
        f'<w:lvlOverride w:ilvl="{i}"><w:startOverride w:val="1"/></w:lvlOverride>'
        for i in range(6)
    )
    extra_nums += (
        f'<w:num w:numId="{nid}">'
        f'<w:abstractNumId w:val="1"/>'
        f'{overrides}'
        f'</w:num>'
    )
num_xml = num_xml.replace('</w:numbering>', extra_nums + '</w:numbering>')
files['word/numbering.xml'] = num_xml.encode('utf-8')

with zipfile.ZipFile(TMP, 'w', zipfile.ZIP_DEFLATED) as zout:
    for n, d in files.items():
        zout.writestr(n, d)

doc = Document(TMP)

# Clear all body content (preserve sectPr)
body = doc.element.body
sectPr = body.find(qn('w:sectPr'))
for child in list(body):
    if child.tag != qn('w:sectPr'):
        body.remove(child)


# Helper: add paragraph with optional numId override
def add_p(style_name, segments, numId=None, ilvl=None,
          align=None, font_size=None, bold_all=False, color=None):
    p = doc.add_paragraph(style=style_name)
    if align is not None:
        p.alignment = align

    # If numId override needed, set numPr on paragraph
    if numId is not None:
        pPr = p._element.get_or_add_pPr()
        # Remove existing numPr if present
        for old in pPr.findall(qn('w:numPr')):
            pPr.remove(old)
        numPr = OxmlElement('w:numPr')
        if ilvl is not None:
            ilvl_el = OxmlElement('w:ilvl')
            ilvl_el.set(qn('w:val'), str(ilvl))
            numPr.append(ilvl_el)
        nid_el = OxmlElement('w:numId')
        nid_el.set(qn('w:val'), str(numId))
        numPr.append(nid_el)
        pPr.append(numPr)

    # Add runs
    if isinstance(segments, str):
        segments = [(segments, False)]
    for text, b in segments:
        run = p.add_run(text)
        if b or bold_all:
            run.bold = True
        if font_size:
            run.font.size = Pt(font_size)
        if color:
            run.font.color.rgb = color
    return p


# Counter for numId per topic
class TopicCtx:
    def __init__(self):
        self.current_num = 100
        self.first_h1 = True

    def new_topic(self):
        self.current_num = 100 + (self.current_num - 100 + 1) % 6
        self.first_h1 = True


tc = TopicCtx()


def H1(text):
    """一、 — first H1 of a topic resets numbering via override."""
    if tc.first_h1:
        p = add_p('Heading 1', text, numId=tc.current_num, ilvl=0)
        tc.first_h1 = False
    else:
        p = add_p('Heading 1', text, numId=tc.current_num, ilvl=0)
    return p


def H2(text):
    return add_p('Heading 2', text, numId=tc.current_num, ilvl=1)


def H3(text):
    return add_p('Heading 3', text, numId=tc.current_num, ilvl=2)


def H4(text):
    return add_p('Heading 4', text, numId=tc.current_num, ilvl=3)


def H5(text):
    return add_p('Heading 5', text, numId=tc.current_num, ilvl=4)


def H6(text):
    return add_p('Heading 6', text, numId=tc.current_num, ilvl=5)


def page_break():
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_topic_title(text, author=''):
    """Topic header: 字型14、粗體 + tab + 作者"""
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(14)
    if author:
        r2 = p.add_run('\t\t\t\t\t' + author)
        r2.bold = True
        r2.font.size = Pt(12)
    # Reset topic counter
    tc.current_num += 1
    if tc.current_num > 105:
        tc.current_num = 100
    tc.first_h1 = True


# ============================================================
# 目錄
# ============================================================
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

p = doc.add_paragraph()
r = p.add_run('目錄')
r.bold = True
r.font.size = Pt(20)

p = doc.add_paragraph()
r = p.add_run('Part 7  Cardiovascular Disease')
r.bold = True
r.font.size = Pt(16)
pPr = p._element.get_or_add_pPr()
shd = OxmlElement('w:shd')
shd.set(qn('w:val'), 'clear')
shd.set(qn('w:color'), 'auto')
shd.set(qn('w:fill'), 'D9D9D9')
pPr.append(shd)

for t in [
    'Topic 56  Venous Thromboembolism Including Pulmonary Embolism',
    'Topic 57  Systemic Hypertension',
    'Topic 58  Pulmonary Hypertension',
    'Topic 59  Aortic Dissection and Related Aortic Syndromes',
    'Topic 60  Aneurysmal Disease',
    'Topic 61  Arterial Occlusion',
]:
    p = doc.add_paragraph()
    r = p.add_run(t)
    r.font.size = Pt(12)

page_break()

# Part header
p = doc.add_paragraph()
r = p.add_run('Part 7  Cardiovascular Disease')
r.bold = True
r.font.size = Pt(16)
pPr = p._element.get_or_add_pPr()
shd = OxmlElement('w:shd')
shd.set(qn('w:val'), 'clear')
shd.set(qn('w:color'), 'auto')
shd.set(qn('w:fill'), 'D9D9D9')
pPr.append(shd)

doc.add_paragraph()

# ============================================================
# Topic 56
# ============================================================
add_topic_title('Topic 56  Venous Thromboembolism Including Pulmonary Embolism', '住院醫師整理')

# 一、簡介與流行病學
H1('簡介與流行病學')
H2('定義')
H3([('肺栓塞 (', False), ('Pulmonary Embolism, PE', True), (' ) 為血栓進入肺動脈循環所致。', False)])
H3([('多源於下肢、上肢、骨盆之 ', False), ('深部靜脈栓塞 (DVT)', True), (' ;偶可源自頸靜脈或下腔靜脈。', False)])
H3([('靜脈血栓栓塞症 (', False), ('VTE', True), (' ) = ', False), ('PE + DVT', True)])

H2('流行病學')
H3([('美國每年新發或復發 PE 約 ', False), ('200,000', True), (' 例;DVT 為其 2 倍。', False)])
H3([('兒童(<18 歲) 急診接受 VTE 檢測者約 1/500,其中 10% 確診;發生率上升,可能與慢性共病存活延長與含 estrogen 口服避孕藥使用增加有關。', False)])
H3([('VTE 發生率隨年齡上升,80 歲達 ', False), ('1/100/年', True)])
H3([('依屍檢資料,PE 為門診病人 ', False), ('第二大突發、非外傷性死因', True)])

H2('預後')
H3([('案例致死率取決血流動力嚴重度、年齡、共病。', False)])
H3([('伴循環性休克者 ', False), ('45%', True), (' ;惟僅 4–5% PE 病人發生休克。', False)])
H3([('血流動力穩定 + <50 歲 + 無共病者致死率約 ', False), ('1%', True)])
H3([('Post-PE syndrome', True), (' :約 25% PE 病人;特徵為慢性疲倦、呼吸困難、運動耐受不良。', False)])
H3([('慢性血栓栓塞性肺高壓', True), (' :約 3%,致殘性呼吸困難。', False)])
H3([('Postthrombotic syndrome', True), (' :約 20% 急診近端 DVT 病人,慢性腿腫與痛。', False)])

# 二、病理生理
H1('病理生理 (Pathophysiology)')
H2([('VTE 分類', False)])
H3([('血栓形成 = 凝血超過纖溶清除之失衡;', False), ('Thrombophilia', True), (' 使凝血優勢化。', False)])
H3([('VTE 分 ', False), ('provoked (誘發)', True), (' 與 ', False), ('unprovoked (idiopathic)', True), (' 兩類。', False)])
H3([('Provoked:近期手術、外傷、肢體或全身固定、活動性癌症、感染、慢性病、estrogen、懷孕/產後、年齡 >50 歲。', False)])
H3([('Unprovoked 1 年復發率 ', False), ('15%', True), (' ;Provoked 為 ', False), ('5%', True), (' (但 Provoked 1 年死亡率較高,多因癌症)。', False)])

H2('血栓位置與肺動脈影響')
H3('臨床顯著之 PE 來源:popliteal、common femoral、superficial femoral、骨盆、axillary、jugular 與大靜脈。')
H3([('至少 ', False), ('1/3 DVT 病人合併無症狀 PE', True)])
H3([('住院 PE 病人 75–80% 影像有 DVT;急診門診 PE 僅 ', False), ('40%', True), (' 合併 DVT。', False)])
H3([('無心肺病人 ', False), ('≥20% 肺血管阻塞', True), (' 才有症狀。', False)])
H3([('clot burden 增加 → 肺動脈壓 ↑ → ', False), ('右心室擴張', True), (' 與心肌損傷 → 釋放 ', False), ('troponin、BNP', True)])
H3([('CT 之 RV/LV ratio 增加、echo 看到右室擴張 + Strain → ', False), ('循環性休克與死亡風險增加', True)])

H2('PE 死亡兩大機轉')
H3('突發近全肺動脈阻塞 → His-Purkinje 缺血 → PEA 或心搏停止。')
H3('進行性右心衰竭與循環性休克 (數小時至數日)。')

H2('DVT 後遺症')
H3([('股、髂股靜脈血栓常於瓣膜形成 → 瓣膜損傷 → 靜脈逆流、鬱積 → ', False), ('postthrombotic syndrome', True), ('(腫、痛、色素沉著、潰瘍)。', False)])

H2('主要危險因子')
H3([('年齡', True), (' :50 歲後上升,80 歲前累積。', False)])
H3([('肥胖', True), (' :BMI >35 起風險上升。', False)])
H3([('懷孕與產後', True), ( ':任一孕期皆可發生。', False)])
H3([('既往 VTE', True), (' :男性、unprovoked、D-dimer 持續上升者復發風險最高。', False)])
H3([('實體癌', True), ( ':腺癌與轉移性疾病風險最大。', False)])
H3([('血液癌', True), (' :急性白血病、myeloma 風險最大。', False)])
H3([('Thrombophilias', True), ( ':非 O 型血、lupus anticoagulant、aPTT 縮短、Factor V Leiden、protein C/S/antithrombin 缺乏。', False)])
H3([('近期手術或大外傷', True), (' :氣管插管或硬膜外麻醉,風險持續至少 ', False), ('4 週', True)])
H3([('固定不動', True), ( ':兩相鄰關節急性固定風險最高(髖+膝)。', False)])
H3([('臥床', True), ( ':約 ', False), ('72 小時', True), (' 後成為危險因子。', False)])
H3([('留置導管', True), ( ':約占上肢 DVT 的 ', False), ('一半', True)])
H3([('長途旅行', True), ( ':一般 ', False), ('>6 小時', True), (' 連續旅行有意義。', False)])
H3([('Estrogen', True), ( ':前幾月最高;含 estrogen 製劑(包括經皮、經陰道)皆增加風險。', False)])
H3([('其他:CHF (依收縮功能不全嚴重度)、stroke (前 1 月最高)、IBD/lupus/nephrotic syndrome。', False)])
H3([('抽菸不獨立增加急診症狀病人之 VTE 機率', True), (' ,但與肥胖、避孕藥協同。', False)])

# 三、PE 臨床特徵
H1('PE 之臨床特徵')
H2('症狀')
H3([('PE 標誌:', False), ('呼吸困難無法以聽診、ECG 變化或 CXR 替代診斷解釋', True)])
H3([('胸痛具肋膜性質為第二常見;但 ', False), ('約一半急診 PE 病人無胸痛', True)])
H3([('Pulmonary infarction 可造成嚴重局部痛(', False), ('embolic event 後 1–3 日', True), (' );基底段可放射至兩肩、類似膽絞痛或輸尿管絞痛。', False)])
H3([('3–4% 急診 PE 有 ', False), ('syncope', True), ( ';1–2% 有 ', False), ('seizure', True), (' 或意識混亂。', False)])
H3([('20% 人有 PFO → ', False), ('paradoxical embolism syndrome', True), (' (中風樣症狀);PFO 預後較差。', False)])

H2('理學檢查')
H3([('生命徵象異常:tachycardia、tachypnea、SpO₂ ↓、輕度發燒;但 ', False), ('PE 無法可靠地改變任何生命徵象', True)])
H3([('約一半 PE 病人 HR <100;約 1/3 早期生命徵象異常會在急診正常化。', False)])
H3('發燒不增加亦不降低 PE 機率。')
H3('兩項有意義的理學發現:')
H4([('單側肢體腫脹', True), ( ' (含留置導管):PE 機率 ', False), ('×3', True)])
H4([('喘鳴', True), ( ':PE 機率 ', False), ('降一半', True)])
H3('多數 PE 聽診清晰;喘鳴或雙側 rales 提示支氣管痙攣或肺炎,但不能完全排除 PE。')
H3('臨床整體 sick/not sick 印象可能誤導(一研究顯示 PE 病人在檢查時更常微笑)。')

H2('Pulse oximetry / CXR / ECG')
H3([('SpO₂ 平均較未確診者低,但 ', False), ('100% 也不能可靠排除 PE', True)])
H3('PE 病人 EtCO₂ 較健康者低。')
H3('CXR 多有非特異異常:cardiomegaly、basilar atelectasis、infiltrate、pleural effusion。')
H3([('<5% 出現 ', False), ('Westermark sign', True), ( ' (肺寡血) 或 ', False), ('Hampton hump', True), (' (肺梗塞圓頂)。', False)])
H3([('ECG 急性肺高壓徵象:HR >100、', False), ('V1–V4 T-wave inversion', True), ( '、incomplete/complete RBBB、', False), ('S1-Q3-T3 pattern', True)])

# 四、DVT
H1('DVT 之臨床特徵')
H2('症狀與徵象')
H3('肢體痛、腫、抽筋。')
H3([('右左腿在脛骨粗隆下 10 cm 直徑差 ', False), ('≥2 cm', True), ( ' 使 DVT 機率倍增。', False)])
H3('上肢導管相關 DVT 常表現為手腫或戒指變緊。')
H3('約 1/4 DVT 有壓痛與紅熱,類似蜂窩組織炎。')
H3([('小腿/saphenous 靜脈血栓較易引發 ', False), ('thrombophlebitis', True)])
H3([('Homan sign', True), ( ' (足背屈引發小腿痛):敏感性與特異性低,', False), ('無預測價值', True)])

H2('威脅肢體之表現')
H3([('近端 DVT 完全阻塞 → 高 compartment 壓:', False)])
H4([('Phlegmasia alba dolens', True), ( ' :蒼白、痛、腫。', False)])
H4([('Phlegmasia cerulea dolens', True), ( ' :青藍、痛、腫。', False)])
H3([('需積極治療包含 ', False), ('clot disruption', True)])

# 五、診斷
H1('VTE 之診斷')
H2('Decision rules 與臨床評估')
H3('估算 pretest probability 為首要步驟;單一檢查或演算法皆無法完美排除/診斷 VTE。')
H3('過度檢查可造成傷害:抗凝出血、誤診後續檢查、對比劑過敏、腎損傷、輻射致癌。')
H3([('Pretest probability ', False), ('<2.5%', True), ( ' 者進一步檢查可能弊大於利。', False)])

H2('PE Rule-Out Criteria (PERC)')
H3([('在臨床 ', False), ('低度懷疑 (gestalt low)', True), ( ' 病人,可可靠地將 PE 機率降至 ', False), ('<2%', True)])
H3('ACEP、ACP、ASH 皆建議使用。')
H3([('關鍵流程:', False), ('先做 gestalt;若低,則用 PERC 決定是否做檢查', True)])
H3('九項條件(全部符合才可排除 PE):')
H4('臨床低機率 (gestalt <15%)')
H4('年齡 <50 歲')
H4('急診全程 HR <100')
H4('SpO₂ >94% (海平面附近;海拔 5000 ft >92%)')
H4('無咯血')
H4('無 VTE 病史')
H4('4 週內無需要氣管插管或硬膜外麻醉的手術/外傷')
H4('無 estrogen 使用')
H4('無單側腿腫(腳跟抬離床面視診不對稱)')

H2('其他預測工具')
H3([('Wells\' rule', True), ( ' :低/中/高機率;具主觀成份(替代診斷可能性)。', False)])
H3([('Modified Wells', True), ( ' :score ≤4 為 low risk,>4 為 non-low risk。', False)])
H3([('Charlotte rule、simplified, revised Geneva (sRGS)', True), (' :排除主觀評估。', False)])
H3([('Unstructured gestalt:低 (<15%)、中 (15–40%)、高 (>40%)。', False)])

H2('D-dimer')
H3('唯一可用以排除 VTE 的血液檢查;原理為 plasmin 分解 fibrin。')
H3([('表示單位:fibrinogen equivalent units (FEU) 或 D-dimer units;1 D-dimer unit = 2 FEU。', False)])
H3([('自動定量法敏感度 ', False), ('94–98%', True), ( ' ,特異度 ', False), ('50–60%', True)])
H3([('半衰期約 ', False), ('8 小時', True), ( ' ,可在症狀性 VTE 後上升至少 3 天。', False)])
H3([('年齡校正:', False), ('age × 10 ng/mL', True), ( '(80 歲 → 800 ng/mL),搭配 Wells ≤4 或 sRGS <5,假陰性僅 0.3%。', False)])
H3([('YEARS 研究:低 pretest + D-dimer 加倍 (例 1000 ng/mL) → PE <1%。', False)])

H2('影像')
H3([('CT pulmonary angiography (CTPA) — PE 影像首選', True)])
H4([('敏感度與特異度約 ', False), ('90%', True)])
H4('高 pretest probability 不能用 CT 排除 PE。')
H4('需平躺、屏氣、IV 導管(20 G 以上)注射約 120 mL 對比劑;不可用中央導管。')
H4('可同時偵測替代診斷(如肺炎 8–22%)。')
H4('subsegmental 充盈缺損 inter-observer agreement 差。')
H4('約 10% CT 因動作偽影或肺動脈不顯影不充分(肥胖、tachypnea)。')
H4('約 15% 出現輕度對比劑腎病變;目前無明確預防措施可有效降低。')
H4('其他併發症:過敏 (<1:1000)、對比劑外滲、繼發 thrombophlebitis。')

H3([('Planar V/Q scan', True)])
H4([('灌流影像同質 → 排除 PE,接近 100% sensitivity', True)])
H4([('兩個或以上 apex central wedge-shaped 缺損 + 通氣正常 → PE 機率 >80%。', False)])
H4('其他結果為 nondiagnostic,單獨無法診斷或排除 PE。')
H4([('V/Q SPECT', True), ( ' :3D、診斷準確度更高;pooled sensitivity 96%/specificity 97%。', False)])

H3([('MRI', True)])
H4('無輻射,孕婦可考慮;但 pooled 敏感度僅 75%、特異度 80%,不建議第一線。')

H3([('Lung US', True)])
H4([('輔助 PE 評估;sensitivity 85%、specificity 83%。', False)])

H3([('Venous US — DVT 影像首選', True)])
H4('迅速、無輻射、直接視覺證據(靜脈不可壓縮)。')
H4('經驗操作者 sensitivity 96%、specificity 96% (DVT);對 PE 為 surrogate 約 40%。')
H4('急診醫師訓練後敏感度 96.1%、特異度 96.8%;兩點壓迫法可能漏診,需至少 10 例經驗。')

# 六、整合性流程
H1('整合性診斷與治療流程')
H2('Step 1:確認生理表現')
H3('至少有一項 PE 之生理表現(症狀如胸/軀幹痛、呼吸問題、意識變化;徵象如生命徵象異常)才開始 PE 檢查。')
H3('未治療氣喘 + 喘鳴 → 先處理支氣管痙攣再評估。')
H3('單純有風險因子但無症狀/PERC 陰性 + gestalt <2% 不需檢查。')

H2('Step 2:是否高於 low suspicion?')
H3('若 gestalt <15% (low) → 用 PERC 排除;PERC 八項全符合即不需進一步檢查。')
H3('若 gestalt 非 low 或 PERC 任一不符 → 安排診斷檢查。')

H2('Step 3:選擇可達 posttest probability <2% 的檢查')
H3('Age-adjusted 定量 D-dimer 為 low/moderate suspicion 病人最佳。')
H3('低 pretest probability + 合理替代診斷、無咯血、無 DVT 徵象 → 可用 D-dimer 兩倍門檻排除。')

H2('Step 4:D-dimer 陽性或高 pretest 才進階檢查')
H3([('優先 ', False), ('CTPA', True), ( ' ;孕婦或腎功能差/對比劑過敏者考慮 ', False), ('V/Q', True)])
H3('Lower extremity venous US:無輻射;敏感度 <40%,US 陰性仍需肺血管影像。')
H3([('Anticoagulation 啟動條件:', False), ('CT 看到節段以上充盈缺損、high probability V/Q、或 US 看到 DVT', True)])
H3('無影像 DVT/PE + D-dimer 陽性 → 可考慮再影像;若有出血風險,2–7 日後重複 venous US。')

H2('DVT 演算法')
H3('低 pretest (<1) 或 modified Wells DVT ≤1 + D-dimer 陰性 → 排除 DVT。')
H3('中高 Wells 或 DVT likely → 直接 US;若 US 陰性 + D-dimer 陽性 → 2–7 日後再 US。')

# 七、治療
H1('VTE 治療')
H2('PE / DVT / 上下肢血栓共通原則')
H3('需 systemic anticoagulation 預防進一步血栓並讓內生纖溶進行。')
H3([('多數初始用 ', False), ('heparin 或 heparin-like', True), ( ' 藥物。', False)])
H3([('急診診斷 DVT 通常於首劑 ', False), ('LMWH、apixaban 或 rivaroxaban', True), ( ' 後出院。', False)])
H3('住院指徵:社會因素、共病、髂股 DVT 合併 phlegmasia。')

H2('抗凝藥物選擇')
H3([('Unfractionated heparin (UFH)', True)])
H4([('80 U/kg bolus → 18 U/kg/h infusion', True)])
H4('適用無法門診治療或嚴重腎衰。')

H3([('LMWH (建議優先)', True)])
H4('Dalteparin 100 IU/kg SC q12h 或 200 IU/kg SC qd')
H4('Enoxaparin 1 mg/kg SC q12h 或 1.5 mg/kg SC qd')
H4('Tinzaparin 175 IU/kg SC qd')

H3([('Factor Xa inhibitor — Fondaparinux', True)])
H4('<50 kg → 5 mg;50–100 kg → 7.5 mg;>100 kg → 10 mg SC qd')
H4([('腎衰禁用', True)])

H3([('Target-specific anticoagulants', True)])
H4('Rivaroxaban 15 mg BID × 21 d → 20 mg qd 隨食;無需 heparin 起導,適合門診。')
H4('Apixaban 10 mg BID × 7 d → 5 mg BID;無需 heparin 起導,適合門診。')
H4([('Dabigatran 150 mg BID;需 ', False), ('heparin run-in 5–10 d', True), ( ' ;腎排泄。', False)])

H2('治療要點')
H3('LMWH 對 PE/DVT 之 composite outcome 與成本優於 UFH(差距不大)。')
H3('PE 不確定但 pretest >20% → empiric 24 h anticoagulation 效益超過風險。')
H3('Heparin 延遲給予增加死亡率;早給但影像前不一定改善 morbidity/mortality。')
H3([('嚴重腎衰 + DVT/PE 多數專家建議 ', False), ('UFH 優於 LMWH', True)])
H3('上肢 DVT 治療同下肢;考慮移除相關留置導管。')
H3([('不要為了 thrombophilia 檢查延誤 UFH', True)])

H2('Phlegmasia cerulea dolens 處置')
H3([('需 ', False), ('迅速降低靜脈壓力', True), ( ' 、肢體置中性、移除壓迫衣物/石膏/敷料。', False)])
H3([('安排血管會診 + ', False), ('catheter-directed thrombolysis', True)])
H3([('若 6 小時內無法轉出且無禁忌 → systemic fibrinolytics(', False), ('Alteplase 50–100 mg IV × 4 h', True), ( ' )。', False)])

H2('Superficial thrombophlebitis 治療')
H3('局部:口服 NSAID 或 topical diclofenac 直至症狀緩解;不需 systemic anticoagulation。')
H3('廣泛淺靜脈受累:full-dose anticoagulation。')
H3('可用壓力襪緩解。')

H2('Calf vein thrombosis')
H3('無共識指引;許多人用 3 個月口服 anticoagulation。')
H3('替代:不治療,1 週後 US 看是否進展;或門診 LMWH。')
H3([('有 VTE 病史或風險因子者:', False), ('3 個月 full-dose anticoagulation', True), ( ' (無禁忌時)。', False)])

H2('PE 門診治療')
H3('精選低風險 PE 病人,可在急診給首劑 LMWH/apixaban/rivaroxaban,觀察後出院。')
H3('短期死亡率與出血率低,成本較低,病人偏好。')
H3([('低風險篩選:', False), ('Modified Hestia 或 Simplified PESI', True)])

H3([('Simplified PESI (0=低風險, 1+=高風險)', True), ( ' :', False)])
H4('>80 歲')
H4('癌症')
H4('HF/CLD')
H4('HR >110')
H4('SBP <100')
H4('SpO₂ <90%')

H3([('Modified Hestia (低風險)', True), ( ' :', False)])
H4('SBP >100')
H4('不需溶栓')
H4('無活動出血')
H4('SpO₂ >94% RA')
H4('未已抗凝')
H4('急診 IV narcotic ≤2 dose')
H4('無其他住院理由')
H4('CrCl >30')
H4('未懷孕、無嚴重肝病或 HIT')

H2('Fibrinolysis')
H3([('PE 嚴重度分類', True)])
H4([('Massive PE', True), ( ' :SBP <90 mmHg >15 min、有 HTN 史的 SBP <100、或 baseline SBP 下降 >40%。', False)])
H4([('Submassive PE', True), ( ' :BP 正常或近正常但有心肺壓力證據(RV/LV ↑、hypokinesis、troponin/BNP ↑、持續低氧)。', False)])
H4([('Less severe', True), ( ' :其餘。', False)])

H3([('適用對象', True)])
H4([('低風險 PE ', False), ('不應給', True), ( ' 溶栓。', False)])
H4([('Massive PE', True), ( ' :受益於溶栓。', False)])
H4([('嚴重 submassive PE 可能受益(存活與生活品質),但 ', False), ('出血風險上升', True)])

H3([('Systemic fibrinolysis 條件', True)])
H4('無禁忌 + 任一:心搏停止、低血壓 (SBP <90)、嚴重呼吸衰竭(SpO₂ <90% + WOB ↑)、echo 右心壓力證據、troponin ↑。')
H4([('禁忌:', False), ('顱內疾病、未控制 HTN、近 3 週大手術/外傷、轉移癌', True)])
H4('暈厥導致頭部外傷者 → 治療前 head CT 排除出血。')
H4([('Alteplase', True), ( ' :PE 唯一核准藥物,', False), ('100 mg IV × 2 h', True)])
H4('Heparin (UFH/LMWH) 通常於溶栓輸注後開始;UFH 維持 aPTT <120 秒。')
H4('50 mg alteplase 可降出血風險,療效相似但仍具爭議。')

H3([('Catheter-directed thrombolysis', True)])
H4([('Intermediate-risk PE 提供良好血流動力改善;顱內出血率 ', False), ('<2%', True)])
H4([('Alteplase 劑量低 ', False), ('(~10 mg)', True), ( ' ,出血風險可能更低。', False)])
H4('>65 歲(顱內出血風險最高)為適用對象。')
H4('因介入準備需時間,massive PE 不建議使用。')

H3([('外科 embolectomy', True)])
H4('年輕、近端大型 PE 合併低血壓者可考慮;常因延遲故 mortality ~30%。')

# 八、特殊族群
H1('特殊族群')
H2('懷孕')
H3('臨床評估困難:正常懷孕也有許多類似徵象。')
H3('Wells ≤4 之孕婦 PE <2%;sRGS 顯示梯度 (low 3.6%, intermediate 9%, high 100%)。')
H3([('PERC 在孕婦尚未充分驗證', True), ( ' ,不建議單獨用以排除。', False)])
H3([('D-dimer 在孕婦特異度低', True), ( ' ;第三孕期幾乎都陽性。', False)])
H3('每孕期上調約 50% (1st 750、2nd 1000、3rd 1250 ng/mL),但未管理性研究驗證。')
H3('孕婦 DVT 之 US 準確度與非孕婦相當。')
H3([('CT 與 V/Q 在技術充分時 sensitivity 100%;MRI 在非孕婦敏感度僅 78%,不建議。', False)])
H3([('治療:', False), ('LMWH', True), ( ' ;massive PE → 系統溶栓、catheter-directed 或 VA-ECMO,母胎存活 >80%。', False)])

H2('Isolated subsegmental PE')
H3('單一肺動脈小充盈缺損 (<3 mm) 且無 DVT;放射科醫師判讀常不一致。')
H3('低復發風險(無 unprovoked VTE、活動癌、其他主要因子)者可能不需 anticoagulation;尚無 RCT 驗證。')
H3('作者建議:apixaban 或 rivaroxaban 門診治療,1 個月後 D-dimer 正常即停藥。')

H2('癌症病人 VTE')
H3([('指引建議:', False), ('LMWH 至少 6 個月', True)])
H3('一 RCT 顯示 rivaroxaban 可用,可降低復發但增加出血。')

page_break()

# ============================================================
# Topic 57
# ============================================================
add_topic_title('Topic 57  Systemic Hypertension', '住院醫師整理')

H1('簡介與流行病學')
H2('盛行率')
H3([('高血壓影響美國 ', False), ('40%', True), ( ' 人口;1–6% 急診病人有嚴重高血壓。', False)])
H3([('其中 ', False), ('1/4 至 1/2 有 end-organ 損傷', True)])
H3('危險因子:肥胖、抽菸、年長、缺乏醫療資源、降壓藥順從不佳。')

H2('高血壓分類 (2017 AHA)')
H3('Normal:<120 / <80。')
H3('Elevated:120–129 / <80。')
H3('Stage 1:130–139 或 80–89。')
H3('Stage 2:≥140 或 ≥90。')

H2('Hypertensive crisis')
H3([('SBP >180 或 DBP >120,包含兩種:', False)])
H4([('Hypertensive emergency', True), ( ' :crisis + ', False), ('end-organ damage', True), ( ' (腦、心、主動脈、腎、眼)。', False)])
H4([('Hypertensive urgency', True), ( ' (爭議性詞彙):無急性 end-organ 失能;>180/120。', False)])
H5([('無臨床獲益', True), ( ';', False), ('急遽降壓可能有害', True)])
H5('建議:重新或加強口服降壓 + 門診迅速追蹤,逐步降壓於數日至數週。')

H1('病理生理')
H2('血管自我調節失衡')
H3('長期 HTN 者動脈壁有生化與結構變化,移動 vascular autoregulatory curve。')
H3('機械應力與 endothelial 損傷 → permeability ↑、hyperperfusion of cerebral/cardiac/renal beds。')
H3('凝血/血小板活化 → fibrin 沉積 → arterioles fibrinoid necrosis。')
H3([('臨床表現:hematuria、retinal hemorrhages、exudates。', False)])

H2('其他機制')
H3('Prostaglandins、free radicals、cytokines、增殖因子 → endothelial 損傷、smooth muscle proliferation、thrombosis。')
H3('RAS 活化 → 血管收縮;Pressure natriuresis → 容積耗損 → 進一步血管收縮 → 末端器官缺血。')
H3('事件後 endothelial 失能可持續多年。')

H1('臨床特徵')
H2('血壓測量原則')
H3([('雙臂測量', True), ( ' ,休息中、短間隔、多次測量再開始治療。', False)])
H3([('雙臂血壓差來源:主動脈夾層、coarctation、PVD、神經肌肉病變;', False), ('差 >10–20 mmHg 有意義', True)])
H3([('檢出後治療較高側,後續測量同手臂', True)])
H3([('避免手腕式振盪測量計', True), ( '(數值偏低)。', False)])
H3('至急診時 SBP/DBP 可下降 ~12 mmHg(無治療);無高血壓史也不能排除 emergency(高達 16% 無病史)。')

H2('胸痛 + 嚴重 HTN')
H3([('快速辨識主動脈夾層至關重要', True), ( ' ;與 ACS 治療不同,且 ACS 抗凝可能在主動脈夾層中災難性。', False)])
H3('主動脈夾層:突發、撕裂/撕扯感、放射至肩胛間。')

H2('急性神經症狀 + 嚴重 HTN')
H3('高血壓 + 頭痛 + 神經缺損 → 缺血/出血性中風。')
H3([('Hypertensive encephalopathy', True), ( ' :排除 focal ischemia/出血後之臨床診斷;意識變化、頭痛、嘔吐、抽搐、視覺異常、papilledema。', False)])
H3([('MRI 顯示 reversible posterior edema → ', False), ('PRES (posterior reversible encephalopathy syndrome)', True), ( ' ,預後差。', False)])

H2('急性腎衰 + 周邊水腫 + 嚴重 HTN')
H3('腳水腫、少尿、食慾不振、噁吐、姿勢性變化、意識改變;部分症狀少。')
H3('血清 Cr 升高、urinary sediment 異常確診。')

H2('Preeclampsia / Eclampsia')
H3([('SBP ≥140 或 DBP ≥90 兩次間隔 ≥4 h, 妊娠 ', False), ('20 週後', True)])
H3([('Eclampsia', True), ( ' :preeclampsia + 新發 grand mal seizure 且無其他神經病因。', False)])

H2('Sympathetic crisis + 嚴重 HTN(四種情境)')
H3([('突然停用 clonidine', True), ( '(口服或經皮);β-blocker 同用會 potentiate。', False)])
H3([('Pheochromocytoma', True), ( ' :罕見,5–20% 為惡性;陣發性頭痛、HTN、tachycardia、diaphoresis。', False)])
H3([('Sympathomimetic 藥物', True), ( ' :cocaine、amphetamine、PCP、LSD;伴 tachycardia、diaphoresis、胸痛、意識變化。', False)])
H3([('MAOI + tyramine', True), ( ' :hyperadrenergic state。', False)])
H3([('Autonomic dysfunction', True), ( ' :脊髓/嚴重頭部外傷、spina bifida;血壓可能僅 marginally 升高,', False), ('不能因接近正常排除', True)])

H2('無症狀但嚴重 HTN')
H3('無正式建議;常做 BMP、ECG、CXR、urinalysis,但無症狀者異常率 <6%。')
H3('依主訴、病史、ROS 選擇檢查。')

H1('治療')
H2('一般原則')
H3('Hypertensive emergency:入加護單位連續 BP 與目標器官監測。')
H3([('Parenteral 降壓藥;', False), ('第 1 小時 SBP 下降 ≤25%', True), ( ' ,穩定後 2–6 h 降至 160/100,接 24–48 h 降至正常。', False)])
H3([('避免 cerebral/coronary/renal hypoperfusion 加劇損傷', True)])
H3('更積極降壓:急性主動脈夾層、pheochromocytoma crisis、嚴重 preeclampsia/eclampsia、急性腦出血。')

H2('主動脈夾層')
H3([('目標:第 1 小時 ', False), ('SBP 100–120, HR ≤60', True)])
H3('降低剪力與主動脈壁應力,限制夾層延伸。')
H3('Opioid 控制疼痛降低交感張力。')

H2('急性肺水腫')
H3('多有未控制 HTN、左心室肥厚、僵硬、舒張功能不全;急性 BP ↑ → afterload ↑、venous capacitance ↓。')
H3([('改善 forward flow(降 afterload)優於 diuresis', True)])
H3([('主治療:', False), ('vasodilators (nitrates 為主)', True)])
H3([('Loop diuretic + nesiritide 合用可能惡化腎功能', True)])
H3('收縮功能不全:IV nicardipine 或 clevidipine 增加 SV 與冠流。')

H2('急性心肌梗塞')
H3([('嚴重 HTN + 缺血變化:SL 或 IV nitrate', True)])
H3([('IV β-blocker 僅建議於嚴重高血壓', True), ( ' ;PO β-blocker 為早期照護一部分但降壓不夠快。', False)])

H2('急性 sympathetic crisis')
H3([('Cocaine/amphetamine:首選 ', False), ('IV benzodiazepine', True), ( ' (lorazepam、diazepam) 降低 adrenergic 刺激。', False)])
H3('監測 respiratory depression 與 sedation;若無效 → 加 nitroglycerin 或 phentolamine;CCB 第三線。')
H3([('β-blocker 可造成 unopposed α-blockade', True), ( ' ,惡化冠狀血管收縮、升 BP;若必須用 ', False), ('labetalol', True), ( '(α-blocking 效果)需配合 vasodilator。', False)])
H3([('Pheochromocytoma:首選 ', False), ('IV phentolamine', True), ( ' (無 IV → IM);第二線 clevidipine、nicardipine;', False), ('Phenoxybenzamine 僅術前', True)])
H3('MAOI toxicity:IV benzodiazepine 為主;不足則 phentolamine、nitroglycerin、nitroprusside。')
H3('達標後密切監測 — 高血壓相期常接續低血壓相。')

H2('急性腎衰')
H3([('Fenoldopam、nicardipine、clevidipine', True), ( ' 皆適合(降 SVR 同時保留腎血流);Fenoldopam 改善 natriuresis 與 CrCl。', False)])

H2('Eclampsia / Preeclampsia')
H3('門檻 BP 較其他 emergency 低。')
H3([('安全:', False), ('Hydralazine、Labetalol', True), ( ' ,另選 ', False), ('oral nifedipine', True)])

H2('神經急症')
H3([('Hypertensive encephalopathy 排除 stroke 後,迅速且穩定降壓:', False), ('IV nicardipine、clevidipine、labetalol、fenoldopam', True)])
H3([('避免 nitroglycerin', True), ( '(擴張腦動脈、改變灌流、可能惡化 autoregulation)。', False)])
H3([('SAH 與 ischemic stroke 之理想 BP 目標未明', True), ( ' ,平衡 ischemia 與 rebleeding。', False)])
H3([('SAH', True), ( ' :labetalol、nicardipine、nitroprusside、clevidipine 皆可;Oral ', False), ('nimodipine', True), ( ' 減少 vasospasm;', False), ('Clazosentan', True), ( ' 高劑量類似效果。', False)])
H3([('Intracerebral hemorrhage', True), ( ' :labetalol、nicardipine、esmolol;Enalaprilat 可用但需 0.625 mg test dose;', False), ('SBP >180 → 130–160 mmHg 可能改善預後', True)])
H3([('Ischemic stroke', True), ( ' :適度 HTN 可保灌注;惡化 edema 與 hemorrhagic transformation。Reperfusion therapy 需 BP ≤185/110;rtPA 後 24 h ≤180/105。', False)])

H1('主要 IV 藥物')
H2('β-Blockers')
H3([('Labetalol', True), ( ' :combined α₁ + 非選擇性 β,', False), ('α:β = 1:7', True), ( ' ;onset 2–5 min,duration 2–4 h;保留腎/腦/冠流;', False), ('孕期安全', True)])
H3([('Esmolol', True), ( ' :ultra-short β1 selective;onset <60 s,duration 10–20 min;適合對 β-blocker 副作用敏感者。', False)])

H2('Calcium Channel Blockers')
H3([('Nicardipine', True), ( ' :第二代 dihydropyridine,腦/冠選擇性;onset 5–15 min,duration 1–4 h。', False)])
H3([('Clevidipine', True), ( ' :ultra-short half-life 2–4 min;選擇性擴張 arteriolar 阻力血管;腎肝獨立代謝。', False)])

H2('Vasodilators')
H3([('Hydralazine', True), ( ' :直接動脈擴張;reflex tachycardia;限孕期使用。', False)])
H3([('Nitroglycerin', True), ( ' :主要 venodilator;低劑量擴張容量血管,高劑量擴 arteriole;', False), ('PDE5i 24–48 h 內禁用', True), ( ' ;首選用於 HF 與 ACS。', False)])
H3([('Nitroprusside', True), ( ' :動靜脈擴張(NO 釋放);onset 數秒,duration 1–2 min;', False), ('腎/肝衰、>24–48 h 有 cyanide/thiocyanate 毒性', True), ( ' ;', False), ('要求 invasive monitoring', True)])

H2('其他')
H3([('Phentolamine', True), ( ' :α1+α2 阻斷;適 pheochromocytoma、cocaine/amphetamine 急症、catecholamine 注射外滲。', False)])
H3([('Fenoldopam', True), ( ' :周邊 ', False), ('dopamine-1 agonist', True), ( ' ;onset 5 min;改善 CrCl、urine flow、Na excretion;適 renal hypertensive emergency。', False)])
H3([('Enalaprilat', True), ( ' :唯一 IV ACEi;可用於 HF 或 ACS,', False), ('first-dose hypotension 注意', True), ( ' ;', False), ('孕婦禁用', True)])
H3([('Clonidine', True), ( ' :central α2-agonist;一般無角色;停藥引起的 rebound 比 β-blocker 嚴重。', False)])

H1('無症狀嚴重 HTN 之治療與處置')
H2('原則')
H3('急性治療不能降低短期罹病率/死亡率;但長期風險高 → 出院前啟動口服降壓並安排追蹤合理。')

H2('常用 PO 藥')
H3([('Carvedilol', True), ( ' 6.25 mg PO,onset 30–60 min,duration 7–10 h。', False)])
H3([('Labetalol', True), ( ' 200–400 mg PO,onset 30–120 min。', False)])
H3([('Metoprolol', True), ( ' 50 mg PO。', False)])
H3([('Captopril', True), ( ' 12.5–25 mg PO,onset 15–30 min。', False)])
H3([('Losartan', True), ( ' 50 mg PO。', False)])
H3([('Hydrochlorothiazide', True), ( ' 25 mg PO,首選之一,onset 2 h,duration 24 h。', False)])
H3([('Nifedipine ER', True), ( ' 10 mg PO(僅 preeclampsia)。', False)])
H3([('Clonidine', True), ( ' 0.1–0.2 mg PO(主要用於 rebound HTN)。', False)])

H2('AHA 2017 處置原則')
H3('SBP 130–139 / DBP 80–89 + ASCVD 風險 <10%:lifestyle + 門診追蹤。')
H3('SBP 130–139 / DBP 80–89 + ASCVD 風險 ≥10%:lifestyle + 啟動降壓(理想 2 種),<1 月追蹤。')
H3('SBP 140–179 / DBP 90–109:評估 target organ damage + lifestyle + 啟動降壓(2 種),約 1 週內追蹤。')
H3('SBP ≥180 / DBP ≥110:lifestyle + 啟動降壓,<1 月追蹤。')

H2('特定共病之首選用藥')
H3([('HF', True), ( ' :Diuretic + ACEi → β-blocker。', False)])
H3([('Post-MI', True), ( ' :β-blocker + ACEi/ARB → Aldosterone antagonist。', False)])
H3([('高 CAD 風險', True), ( ' :β-blocker + CCB(若 angina)。', False)])
H3([('再發中風預防', True), ( ' :Thiazide + ACEi/ARB。', False)])
H3([('Diabetes', True), ( ' :Thiazide、ACEi、ARB、CCB。', False)])
H3([('CKD', True), ( ' :ACEi 或 ARB。', False)])
H3('用藥前需 chemistry panel (Cr、K)、ECG、必要時 pregnancy test。')

page_break()

# ============================================================
# Topic 58
# ============================================================
add_topic_title('Topic 58  Pulmonary Hypertension', '住院醫師整理')

H1('簡介與流行病學')
H2('正常生理')
H3('肺血管系統正常為高流量、低阻力;mPAP 約佔系統循環之 15–20%。')
H3('正常 PA systolic 15–30 mmHg;diastolic 4–12 mmHg。')

H2('PH 定義')
H3([('靜止 mPAP >25 mmHg', True), ( ' 或運動時 >30 mmHg。', False)])
H3('依 mPAP、PVR、PCWP 分類;echo 估算,確診仍需 right heart catheterization。')

H2('WHO 分類')
H3([('Group 1 — Pulmonary Arterial Hypertension', True)])
H4([('mPAP >25, PVR >240 dynes·s·cm⁻⁵, PCWP <15。', False)])
H4('Idiopathic、heritable、藥物/毒素、HIV/portal HTN/PVOD/結締組織疾病。')
H3([('Group 2 — Pulmonary Venous Hypertension(left heart disease,', False), ('最常見', True), ( ' )', False)])
H4('左室收縮/舒張功能不全、mitral/aortic 瓣膜疾病。')
H3([('Group 3 — Chronic Hypoxemic Lung Disease', True)])
H4('COPD、ILD、IPF、結締組織疾病、OSA、高海拔暴露。')
H3([('Group 4 — Chronic Thromboembolic PH (CTEPH)', True)])
H4([('VTE 病人 ', False), ('最高 4%', True), ( ' 發生。', False)])
H3([('Group 5 — Miscellaneous', True)])
H4('Lymphatic obstruction、myeloproliferative、sarcoidosis、neurofibromatosis、glycogen storage、甲狀腺。')

H2('預後')
H3([('Idiopathic PAH ', False), ('5 年死亡率 >30%', True)])

H1('病理生理')
H2('血管重塑')
H3('小及中型肺動脈血管重塑、發炎、原位血栓形成。')
H3('微血管 permeability 改變、缺氧性血管收縮、arteriolar plexiform lesion。')
H3('累積 → 持續 PVR 升高 + 肺血流受損。')

H2('右心室衰竭發展')
H3([('PVR 持續升高 → ', False), ('RV dilation', True), ( ' → 收縮期壁張力 ↑、O₂ 消耗 ↑、收縮力 ↓。', False)])
H3('進行性 RV 擴張 → 室間隔向 LV 偏 → LV 充填減少 → CO 與灌注下降。')

H2('右冠灌流')
H3([('右冠灌流 = systemic diastolic − RA pressure', True)])
H3('進階 PH:RCA 灌注幾乎僅在 diastole → RV ischemia → CO 進一步降低 → RV 衰竭、心血管崩潰。')

H1('臨床特徵')
H2('症狀')
H3([('最常見:', False), ('呼吸困難', True), ( '(休息或運動時),>50%。', False)])
H3('其他:疲倦、胸痛、暈厥、運動性 lightheadedness。')
H3([('診斷常被延誤,平均症狀至診斷 ', False), ('2 年', True)])
H3('進階 HF 症狀:早飽、厭食、orthopnea、PND、周邊水腫。')

H2('理學檢查')
H3('早期常正常。')
H3('進展後 RV 衰竭徵象:holosystolic TR murmur、JVD、肝腫大、腹水、P2 增強、下肢水腫。')

H1('診斷檢查')
H2('ECG')
H3([('最常見:', False), ('right axis deviation', True)])
H3('其他:V1 R/S >1、V5/V6 R/S <1、qR in V1、S1Q3T3、inferior 之 RA enlargement、incomplete/complete RBBB。')
H3([('敏感度與特異度均不高', True)])
H3([('可顯示 RV ischemia 或心律不整;最常見:', False), ('Afib、Aflutter、AVNRT', True)])

H2('實驗室')
H3('CBC、CMP 多非特異。')
H3('Troponin/BNP 上升 → 較嚴重疾病。')
H3('肝功能、lactate、凝血異常 → 肝鬱血,預後不良。')

H2('影像')
H3('CXR 異常:RA、RV、肺門 PA 擴大;依病因可有肺水腫、過度充氣、ILD。')
H3([('TTE 為急診最佳初評工具', True), ( ' :', False)])
H4('RV 肥厚、RV 功能下降、TR。')
H4('偵測 RV 衰竭誘因(節段性壁運動異常、瓣膜問題)。')
H4([('Severe 時 IVS leftward deviation、', False), ('RV/LV end-diastolic >1', True), ( '(four-chamber)。', False)])
H3([('US 證據:dilated IVC + hepatic vein、parasternal short-axis 之 D-shaped septum、A4C 之 RV 肥厚與 LV 受壓。', False)])
H3([('CTPA 為 PE 首選', True), ( ' ;不可用 TTE 排除 VTE;PH 病人 VTE 風險高 → 早期 CT 並評估 RV。', False)])

H1('治療')
H2('Oxygen 與機械通氣')
H3([('共識:', False), ('SpO₂ >90%', True), ( '(不必常規追求更高)。', False)])
H3([('插管與正壓通氣常 ', False), ('惡化 RV 衰竭', True), ( ' (intrathoracic pressure ↑、coronary perfusion ↓)→ 心血管崩潰。', False)])
H3([('需通氣:lung-protective(', False), ('TV 6–8 mL/kg IBW、最低 PEEP', True), ( ' )、plateau pressure <30 cmH₂O。', False)])
H3([('避免 hypercapnia', True), ( '(升 PVR、PA pressure、RV strain)。', False)])

H2('容積管理')
H3('容積過載 → RV 擴張 → LV 輸出受損。')
H3([('低血容量:', False), ('100–250 mL', True), ( ' 等張晶體挑戰。', False)])
H3([('右側壓力高 → ', False), ('動態容積反應指標可能不可靠', True)])
H3('最終可考慮 diuresis 移除多餘容積。')

H2('RV function')
H3('無低血壓之 RV 衰竭:inotropic 增加收縮力與降 RV afterload。')
H3([('Dobutamine (β1 agonist) 為首選', True), ( ' :', False), ('2 → 10 µg/kg/min', True), ( ' ;', False), ('>10 µg/kg/min', True), ( ' 可致 tachydysrhythmia 與 hypotension。', False)])
H3([('不耐 dobutamine:', False), ('Milrinone', True), ( ' (PDE-3 inhibitor) 0.125 → 0.75 µg/kg/min;高劑量可致 hypotension。', False)])

H2('Right coronary artery perfusion')
H3('為維持 RV 功能,主動脈根 BP 必須高於 PA pressure。')
H3([('低血壓 PH 病人:', False), ('Norepinephrine', True), ( ' 0.05 µg/kg/min 起始;避免高劑量。', False)])
H3([('避免 dopamine (tachydysrhythmias) 與 phenylephrine (益處較少)。', False)])

H2('RV afterload')
H3('門診慢性 PH 常用 pulmonary vasodilator:prostanoids、ERA、PDE5i。')
H3([('IV prostanoid 居家治療之病人:', False), ('立刻重啟 home infusion', True), ( ' ,確認導管與幫浦;若阻塞 → 立即改 peripheral IV 給藥。', False)])
H3([('急性發作合併 RV 衰竭可給 ', False), ('inhaled epoprostenol 或 inhaled NO', True), ( '(facemask、HFNC、ETT) 降 RV afterload。', False)])
H3([('急診不建議啟動 IV pulmonary vasodilator', True)])

H1('常用門診 pulmonary vasodilator')
H2('Prostanoids')
H3([('Epoprostenol', True), ( ' :2–5 min 半衰期,連續 IV;急性 PH RV 衰竭可氣霧劑。', False)])
H3([('Treprostinil', True), ( ' :IV/SC,4–5 h 半衰期;', False), ('不適 acute RV failure', True)])
H3([('Iloprost', True), ( ' :氣霧劑;不耐 parenteral 者使用;不適 acute RV failure。', False)])

H2('Endothelin Receptor Antagonist (ERA)')
H3([('Bosentan、Ambrisentan', True), ( ' :PO;不適 acute RV failure;副作用 transaminase ↑、Hb ↓。', False)])

H2('Phosphodiesterase-5 Inhibitor')
H3([('Sildenafil、Tadalafil', True), ( ' :PO;不適 acute RV failure;副作用:頭痛、flushing、dyspepsia、與 nitrate 合用低血壓。', False)])

page_break()

# ============================================================
# Topic 59
# ============================================================
add_topic_title('Topic 59  Aortic Dissection and Related Aortic Syndromes', '住院醫師整理')

H1('簡介與流行病學')
H2('範圍')
H3([('急性主動脈症候群', True), ( ' 包括:', False), ('主動脈剝離、PAU、IMH、主動脈瘤滲漏與破裂', True)])
H3([('發生率:', False), ('每 10 萬人 2.9–4.7 例/年', True)])
H3([('1/5/10 年存活(手術病人):', False), ('92% / 77% / 57%', True)])
H3([('22% 死前未診斷', True)])

H2('結締組織疾病')
H3([('Marfan 最常見心血管併發症為 ', False), ('aortic root disease 與 type A dissection', True)])
H3([('辨識 ', False), ('TGFBR2、FBN1', True), ( ' 等基因突變 + 規律追蹤可降低致命結果。', False)])

H1('病理生理')
H2('共同機轉')
H3('發生於慢性高血壓與其他造成主動脈中膜退化之背景。')
H3([('易感:', False), ('bicuspid aortic valve、Marfan、Ehlers-Danlos、家族史', True)])
H3('慢性 cocaine/amphetamine 加速 atherosclerosis 並升風險;既往心臟手術為危險因子。')
H3('所有機轉:中膜層削弱 + 內膜壁應力增加。')

H2('Aortic dissection')
H3('內膜破裂後血液進入中膜並於內、外膜間剝離。')
H3('兩大撕裂處:')
H4([('Sinotubular junction', True), ( '(升主動脈起點)', False), ( ' 50–65%', True)])
H4([('左鎖骨下動脈遠端', True), ( ' (升降主動脈交界)', False), ( ' 20–30%', True)])
H3('剝離血柱形成 false lumen,可向遠端(最常)、近端、雙向延伸。')
H3([('血液若再穿入內膜重建血流,', False), ('可能誤導為自發性緩解', True)])
H3([('穿透外膜 → ', False), ('幾乎都迅速致命', True)])
H3('雙峰年齡分布:年輕(結締組織疾病)/ >50 歲(慢性 HTN/IHD)。')
H3('既往剝離為再剝離的危險因子。')

H2('分類系統')
H3([('Stanford', True), ( ' :', False)])
H4([('Type A', True), ( ' :任何含升主動脈。', False)])
H4([('Type B', True), ( ' :僅降主動脈。', False)])
H3([('DeBakey', True), ( ' :', False)])
H4([('Type 1', True), ( ' :升 + 弓 + 降。', False)])
H4([('Type 2', True), ( ' :僅升。', False)])
H4([('Type 3', True), ( ' :僅降。', False)])

H2('IMH 與 PAU')
H3([('Intramural hematoma (IMH)', True), ( ' :中膜梗塞,通常因 vasa vasorum 損傷;可自發緩解或進展為剝離。', False)])
H3([('Penetrating atherosclerotic ulcer (PAU)', True), ( ' :可導致 IMH、剝離或穿孔。', False)])
H3([('IMH 約 ', False), ('62% 為女性', True), ( ' ;女性剝離高齡發病、住院併發症多、院內死亡率較高。', False)])

H1('臨床特徵')
H2('病史')
H3('初始撕裂部位常決定症狀;隨延伸/牽涉其他動脈或器官變化。')
H3([('典型:', False), ('突發', True), ( ' 、', False), ('劇烈胸痛', True), ( ' 、放射至 ', False), ('肩胛間', True), ( ' ,伴瀕死感。', False)])
H3('464 例研究:60% 前胸痛(Type A 較多);腹痛較多見 Type B。')
H3([('描述:64% sharp;50% tearing/ripping;', False), ('40% 無胸痛(亞洲研究)', True)])
H3([('Syncope ~10%', True), ( '(Type A 較常,提示心包填塞)。', False)])
H3('22% 剝離發生於有心臟手術史病人。')
H3([('近頸動脈剝離可呈中風;', False), ('20% Type A 有神經學徵候 → 預後較差', True)])
H3([('脊髓血流中斷 → ', False), ('截癱', True)])
H3('進一步遠端 → 背、側腹、腹痛。')
H3([('近端至主動脈根 → ', False), ('心包填塞', True), ( ' ,通常致命。', False)])

H2('理學檢查')
H3([('多數病人 ', False), ('檢查相對正常', True)])
H3('AR murmur (32%)、橈/股動脈脈搏不對稱 (15%)。')
H3([('雙臂 SBP 差 ', False), ('>20 mmHg', True), ( ' 與剝離獨立相關;但 19% 無剝離者也有。', False)])
H3([('HTN 49%;', False), ('低血壓 18–25%(預後較差)', True)])
H3('主動脈擴張壓迫:食道、recurrent laryngeal nerve、上頸交感神經節 → 吞嚥困難、聲嘶、Horner 症候群。')

H2('Aortic Dissection Detection Risk Score (ADD-RS)')
H3('三類各有任一項給 1 分,總分 0–3。')

H3([('類別 1:背景條件', True)])
H4('Marfan 症候群')
H4('主動脈疾病家族史')
H4('主動脈瓣疾病')
H4('近期主動脈操作')
H4('胸主動脈瘤')

H3([('類別 2:疼痛特徵(胸/背/腹)', True)])
H4('突發')
H4('劇烈')
H4('撕裂/撕扯感')

H3([('類別 3:理學檢查異常', True)])
H4('四肢脈搏振幅或 SBP 差')
H4('Focal 神經缺損 + 胸/背/腹痛')
H4('新發 AR murmur + 痛')
H4('休克或低血壓')

H1('診斷')
H2('鑑別困難')
H3('胸痛 + 多種末端缺血表現 → 鑑別困難。')
H3([('鑑別:', False), ('MI/ACS、心包疾病、stroke、肌肉骨骼、脊髓、腹腔內疾病、PE/肺炎/胸膜炎/氣胸', True)])
H3([('剝離破裂入真腔 → 症狀停 → ', False), ('可能誤撤診斷', True)])
H3('誤診相關:walk-in、CXR 縱膈/主動脈輪廓正常、雙側脈搏對稱、非特異症狀。')

H2('ECG')
H3('Q wave / ST elevation:3–4%。')
H3('ST 下降:15–22%。')
H3('非特異 ST/T:41–62%。')
H3('完全正常:19–31%。')

H2('生物標記:D-dimer')
H3('Cutoff 500 ng/mL,sensitivity 98%,LR− 0.05;specificity 41%。')
H3([('指引 ', False), ('不建議單以 D-dimer 排除剝離', True), ( ';false negative 高達 ', False), ('18%', True)])
H3('ADD-RS ≤1 + D-dimer <500 → 剝離率僅 0.3%(待外部驗證)。')

H2('影像')
H3([('CXR', True), ( ' :', False)])
H4([('12–37% 完全正常 → ', False), ('不能排除', True)])
H4('最常見:縱膈擴大或主動脈輪廓異常。')
H4('其他:胸水、內膜鈣化移位、氣管/主支氣管/食道偏移。')

H3([('CT 為診斷影像首選', True)])
H4('顯示 false lumen、剝離瓣、延伸、破裂徵象、末端器官損傷。')
H4('含對比與不含對比皆做。')

H3([('TEE', True)])
H4('經驗操作者敏感性與 CT 相當。')
H4('通常需鎮靜或全麻;食道病為相對禁忌。')
H4('氣管/左支氣管空氣干擾升主動脈評估。')

H3([('MRI', True), ( ' :用於穩定病人。', False)])

H3([('Triple rule-out (CTA 同時評估 CAD/PE/dissection)', True), ( ' :未證實改善診斷率或臨床結果,', False), ('不推薦', True)])

H1('治療')
H2('降壓:負性肌力藥優先')
H3('剝離可低血壓需復甦,但多數需降壓。')
H3([('初始用 ', False), ('負性肌力藥', True), ( ' 降壓且不增加 shear force。', False)])
H3([('β-blocker 為首選', True), ( ' ;短效 esmolol、labetalol 優於長效。', False)])
H3([('SBP 目標 ', False), ('120–130 mmHg', True), ( ' ;部分指引建議 ', False), ('100–120 mmHg', True), ( '(無 AR)。', False)])
H3([('Esmolol', True), ( ' :0.1–0.5 mg/kg IV bolus/1 min → 0.025–0.2 mg/kg/min infusion。', False)])
H3([('Labetalol', True), ( ' :10–20 mg IV → q10min 20–40 mg(max 300 mg)。', False)])
H3([('β-blocker 與 IRAD 改善存活相關', True)])

H2('Vasodilator')
H3([('成功給予 esmolol/labetalol 後可加 ', False), ('IV nicardipine、clevidipine、nitroglycerin、nitroprusside', True)])

H1('處置與追蹤')
H2('立即處置')
H3([('立即會診 ', False), ('血管外科', True)])
H3('多數仍以 open repair 為主流;EVAR 漸增。')
H3('需 ICU 進行血流動力治療與監測。')

H2('IMH/PAU')
H3('自然病程不明,治療仍具爭議。')

H2('出院原則')
H3([('任何急性主動脈症候群病人 ', False), ('不應未經會診出院', True)])

H1('特殊族群:懷孕合併剝離')
H2('流行病學')
H3([('罕見;多於 ', False), ('第三孕期及產後', True)])
H3('危險因子:bicuspid aortic valve、結締組織病、HTN、家族史。')
H3([('Marfan 孕婦中 ', False), ('4.4% 妊娠合併剝離', True)])

H2('處置')
H3([('診斷一旦考量,需 ', False), ('同時會診產科及心血管外科', True)])

page_break()

# ============================================================
# Topic 60
# ============================================================
add_topic_title('Topic 60  Aneurysmal Disease', '住院醫師整理')

H1('簡介')
H2('定義')
H3([('動脈瘤', True), ( ' :動脈壁擴張至正常的 ', False), ('1.5 倍以上', True)])
H3([('分類:', False), ('true aneurysm、pseudoaneurysm、mycotic aneurysm', True)])
H3([('True aneurysm', True), ( ' :含血管全層。', False)])

H2('危險因子')
H3('抽菸、年長、白人、HTN、高血脂、結締組織病、家族史。')

H1('病理生理')
H2('動脈瘤形成')
H3('彈性蛋白、collagen、fibrolamellar units 漸減 → 中膜變薄、張力下降。')
H3([('Laplace law:', False), ('wall tension = pressure × radius', True), ( ' → 越大越擴張越快。', False)])
H3([('擴張速率多變,大者較快;', False), ('每年平均 0.25–0.5 cm', True), ( ' ,但突發擴張無法預測。', False)])
H3([('破裂為災難性事件', True), ( ' ,壁應力超過張力即發生。', False)])
H3('社經地位低者較常以破裂表現,術後預後較差。')

H2('Pseudoaneurysm')
H3('壁部分為血管,部分為纖維/周邊組織。')
H3('原因:既往導管、吻合術、外傷、感染。')
H3('小者可自發栓塞。')

H2('Mycotic / Infected aneurysm')
H3([('Mycotic', True), ( ' :來自瓣膜內膜炎之敗血性栓塞;免疫低下者好發。', False)])
H3([('Infected', True), ( ' :既存動脈瘤合併菌血症或鄰近感染。', False)])
H3([('IVDU 為兩者主要危險因子', True), ( ' ,', False), ('小心誤診為蜂窩組織炎/膿瘍', True)])

H2('周邊與內臟動脈瘤')
H3([('Popliteal 為最常見周邊動脈瘤', True), ( ' ,常合併對側或 AAA。', False)])
H3('Femoral 較少見,常合併他處動脈瘤。')
H3('內臟 (renal、splenic、hepatic) 多靜默,直至破裂等併發症。')
H3('男性常見(脾動脈瘤除外)。')
H3([('併發症:破裂(', False), ('mortality 80%', True), ( ' )、栓塞致缺血。', False)])

H1('一般臨床特徵')
H2('表現')
H3('臨床表現非特異;依位置、對周邊壓迫、栓塞而異。')
H3('內臟動脈瘤常於腹/側腹影像偶發發現;下肢常於 Doppler 找 DVT 時發現。')
H3([('軀幹動脈瘤一旦破裂 → ', False), ('出血性休克', True), ( ' ,無迅速手術死亡率高。', False)])

H1('有症狀腹主動脈瘤 (AAA)')
H2('定義')
H3([('AAA:腹主動脈直徑 ', False), ('≥3.0 cm', True)])
H3([('修復考量:', False), ('≥5.0 cm', True)])

H2('流行病學')
H3('一等親有 AAA 者 18%;>60 歲、男性風險高。')
H3([('抽菸為最重要環境危險因子', True), ( ' ;抽菸者盛行率為終生不吸菸者 4 倍以上;為加速擴張與破裂之主要因子。', False)])

H2('臨床特徵')
H3([('最常:', False), ('背痛或腹痛', True), ( ' ,突發、劇烈;約半數撕裂/撕扯感。', False)])
H3([('經典三聯徵', True), ( ' :', False), ('腹痛 + 搏動性腫塊 + 低血壓', True), ( ' ,', False), ('僅 1/3 具備', True)])
H3([('Syncope ~10%', True)])
H3('非典型疼痛部位:側腹、鼠蹊、單側象限、髖部。')
H3('其他:噁吐、膀胱痛、髖痛、tenesmus。')
H3([('11% 否認疼痛', True)])
H3([('Aortoenteric fistula', True), ( ' :GI bleeding;肢端缺血(血栓栓塞);休克;猝死。', False)])
H3([('Syncope 無預警 + 隨後劇烈痛 → ', False), ('破裂後暫時局限', True)])

H2('診斷 — 理學檢查')
H3([('觸診依大小:3.0–3.9 cm = ', False), ('29%', True), ( ';4.0–4.9 cm = ', False), ('50%', True), ( ';≥5.0 cm = ', False), ('76%', True)])
H3([('觸壓痛常被解讀為擴大/破裂,但 ', False), ('無壓痛不代表完整', True)])
H3('肥胖難觸診;消瘦者尺寸常被高估。')
H3('低血壓與心搏過速僅少數具備。')

H2('鑑別與常見誤導')
H3('鑑別:syncope、腹/胸/背痛、休克之原因。')
H3('合併 CAD、慢性肺病可分散注意力。')
H3([('小心 ', False), ('誤診為腎結石', True)])
H3([('Cullen sign', True), ( '(臍周瘀斑)、', False), ('Grey Turner sign', True), ( '(側腹瘀斑) 罕見。', False)])
H3('後腹腔血液可至會陰、鼠蹊 → 陰囊/外陰血腫。')
H3('Iliopsoas sign;股神經壓迫致神經病變。')
H3([('AAA 通常 ', False), ('不影響股動脈搏動', True)])

H2('Aortoenteric fistula')
H3('無肝病的不明大量上下消化道出血,主動脈移植史增風險;最常侵犯十二指腸。')
H3('小量哨兵出血可為首發。')
H3('Aortovenous fistula:高輸出心衰、遠端動脈血流降、中央靜脈容積上升。')

H2('影像')
H3([('不穩定病人 ', False), ('不可送出急診', True)])
H3([('床邊 US 為首選', True), ( '(', False), ('>90% 敏感', True), ( ' ),測量 ', False), ('外壁到外壁、橫切+縱切', True), ( ' ;辨識 SMA 區分主動脈與下腔。', False)])
H3([('主動脈 ', False), ('<3.0 cm 可排除', True)])
H3([('CT 含對比劑為穩定病人首選', True), ( ' ;不能對比劑者用無強化 CT。', False)])
H3([('平片偶見 65% 鈣化(側位較明顯),但 ', False), ('不能排除 AAA', True)])

H2('治療')
H3([('關鍵:辨識並 ', False), ('迅速手術控制出血', True)])
H3([('臨床懷疑時立即會診;低血壓或末端灌注不足時 ', False), ('不可延遲影像', True)])
H3([('標準急救:', False), ('兩條大號 IV', True), ( ' 、心電監測、補氧。', False)])
H3([('積極輸液可能有害', True), ( ' ;', False), ('Permissive hypotension', True), ( ' :SBP ', False), ('80–90 mmHg', True), ( ' ,意識為指標。', False)])
H3([('Esmolol(半衰期 9 min):懷疑擴大且嚴重 HTN 時可滴定至 ', False), ('SBP 120 mmHg', True)])
H3([('REBOA', True), ( ' :部分中心使用;研究小且 mortality 改善不明;主用於外傷。', False)])
H3([('任何大小有症狀皆視為緊急', True), ( ' ,無手術能力應啟動復甦並轉院。', False)])
H3([('立即轉院並由能辨識/治療休克者陪伴', True)])
H3([('有症狀者圍術期 ', False), ('mortality ×2', True), ( ' ;破裂者 ', False), ('×7', True)])

H2('後續追蹤')
H3('所有無症狀皆需追蹤。')
H3([('≥5 cm 風險高 → 數日內追蹤', True)])
H3([('女性 AAA 較小尺寸即可能破裂', True), ( ';懷孕及產後風險上升。', False)])
H3('3–5 cm 由初級照護或外科追蹤,血壓控制。')
H3('男性 <5.5 cm 修復未改善存活,女性可能適合。')
H3([('EVAR 為較不侵入選項;併發症:', False), ('graft leak、阻塞、腎功能不全、aortoenteric fistula、腸缺血', True)])

H1('胸主動脈瘤 (TAA)')
H2('表現與處置')
H3('可壓迫或侵蝕食道、氣管、支氣管、神經。')
H3([('侵蝕鄰近結構 ', False), ('通常立即致死', True), ( ' ,除非有迅速復甦。', False)])
H3([('無症狀者轉介外科 + 初級照護,', False), ('血壓控制', True)])
H3('修復可採 open 或 endovascular,長期結果仍具爭議。')

H1('肢體與內臟動脈瘤')
H2('一般原則')
H3('肢體(膕、鎖骨下、股):破裂、血栓、周邊栓塞 → 肢體威脅性。')
H3('內臟動脈瘤通常靜默至併發症;包括 renovascular HTN、腎動脈血栓、遠端栓塞致器官梗塞、AV fistula、破裂。')

H2('膕動脈瘤 (Popliteal)')
H3('定義:>2 cm 或 >正常 150%。')
H3('膕窩不適、合併 DVT 之腿腫、claudication。')
H3([('破裂罕見', True), ( ' ;最嚴重併發症為 ', False), ('血栓/栓塞 → 急性肢體缺血', True)])
H3('觸診:膕窩硬塊或搏動性腫塊。')
H3([('診斷:', False), ('動脈 US,評估雙側', True)])

H2('股動脈/髂動脈瘤')
H3('表現:鼠蹊或大腿上端搏動性腫塊、陰囊血腫、急性肢體缺血。')
H3('髂動脈瘤難以臨床懷疑,因無法直接觸診;症狀像泌尿/腸道/鼠蹊疾患。')

H2('肝動脈瘤')
H3('動脈粥狀硬化或外傷造成;破裂可致急性出血性休克。')
H3([('Quincke triad', True), ( ' :', False), ('黃疸、膽絞痛、上消化道出血', True), ( '(hemobilia)。', False)])

H2('脾動脈瘤')
H3('左上腹痛、未分化休克、腹腔內出血。')
H3('破裂預後差(腹腔位置、表現非特異)。')
H3([('特別注意 ', False), ('懷孕第三孕期', True), ( ' ;門脈高壓亦為危險因子。', False)])

H2('鎖骨下與無名動脈瘤')
H3('上肢缺血表現:疼痛、無力、蒼白、感覺異常。')

H2('吻合動脈瘤')
H3('可見於主動脈、髂、股動脈;破裂可致大出血。')
H3([('可有 ', False), ('哨兵小出血', True), ( '(疼痛或短暫低血壓)。', False)])
H3([('可侵蝕鄰近腸 → ', False), ('Aortoenteric fistula', True), ( ' ;AAA 修復史 + 嚴重下消化道出血為經典表現。', False)])

H2('影像與治療')
H3('Doppler US 為周邊動脈瘤之初步無創檢查。')
H3([('合併缺血/血栓 → 緊急 ', False), ('CT angiography', True), ( ' 兼具診斷與治療規劃。', False)])
H3('有症狀者迅速會診血管外科。')
H3('無症狀周邊動脈瘤門診處置。')
H3([('Mycotic / infected aneurysm:', False), ('長期 IV 抗生素 + 手術', True), ( ';疑似破裂立即會診。', False)])

page_break()

# ============================================================
# Topic 61
# ============================================================
add_topic_title('Topic 61  Arterial Occlusion', '住院醫師整理')

H1('簡介與流行病學')
H2('定義')
H3([('肢體缺血', True), ( ' :血流受限導致組織灌注不足出現症狀。', False)])
H3([('Acute limb ischemia (ALI)', True), ( ' :新發或進展之血流缺損,需立即辨識、迅速治療以挽救肢體。', False)])
H3([('Critical limb ischemia (CLI)', True), ( ' :PAD 譜的末端,出現 ', False), ('靜止痛、潰瘍、壞疽', True)])

H2('危險因子')
H3([('最重要:', False), ('抽菸、糖尿病', True)])
H3('其他:高血脂、HTN、homocysteine ↑、CRP ↑。')

H2('流行病學')
H3('PAD 嚴重度 ↔ MI、ischemic stroke、血管性死亡風險。')
H3([('最常受影響:', False), ('股膕 > 脛 > 主動脈髂 > 頭臂', True)])
H3([('CLI 1 年死亡率 ', False), ('25%', True), ( ' ,', False), ('1/4 存活者需截肢', True)])
H3('全球 PAD >2 億人;美國 CLI 年發生率 0.35%、平均盛行率 1.33%。')

H1('病理生理')
H2('總論')
H3('ALI:急性血供下降 → 組織低灌注 → 威脅肢體存活。')
H3([('無側支循環時,周邊神經與骨骼肌可於 ', False), ('4–6 h', True), ( ' 內出現不可逆變化。', False)])
H3([('血流恢復後可有 ', False), ('再灌流損傷', True), ( ' :compartment syndrome、橫紋肌溶解、代謝紊亂、', False), ('高血鉀、myoglobinemia、代謝性酸中毒、CK ↑', True)])
H3('傷害程度依阻塞時間、位置、側支循環、原本肢體健康。')
H3([('1/3 動脈阻塞死亡來自再血管化後代謝併發症', True)])

H2('血栓 (Thrombosis)')
H3([('最常見原因', True), ( ' ,可見於原生血管或 bypass graft。', False)])
H3([('下肢:', False), ('>80% 為血栓性', True)])
H3('上肢:約 1/2 血栓、約 1/3 栓塞。')
H3('臨床上不易區分血栓與栓塞。')
H3('血栓多見於 atherosclerosis 背景;有側支故發作較漸進。')
H3('斑塊破裂可直接栓塞或誘發栓塞性血栓。')
H3('進展機轉:')
H4('血栓延伸阻塞側支')
H4('缺血性水腫致 compartment syndrome')
H4('血栓碎片進入微循環')
H4('微血管細胞水腫')
H3([('大血管再灌流不一定能解除微循環阻塞', True)])
H3('正常血管罕見發生 → 評估高凝狀態。')

H2('栓塞 (Embolism)')
H3([('心臟為周邊栓塞之主要來源;', False), ('Afib 最常見', True)])
H3([('近期 MI 之 ventricular mural thrombus:24 h 至 3 個月,整體 3%,', False), ('前壁 MI 達 9%', True)])
H3('其他:機械瓣膜、心房黏液瘤、瓣膜贅生物、人工裝置碎片。')
H3('非心源性:動脈瘤與粥樣斑塊之血栓(主動脈髂、股、膕、鎖骨下)。')
H3([('Atheroemboli', True), ( ' :膽固醇與血小板聚合物造成微循環阻塞 → ', False), ('blue toe syndrome', True), ( ' 、手部表現、TIA。', False)])
H3([('Paradoxical embolization', True), ( ' :靜脈血栓經 PFO 至左側循環。', False)])
H3([('栓子常停滯於 ', False), ('狹窄/分叉處', True), ( ';下肢:', False), ('股總動脈分叉處 > 膕動脈', True), ( ';上肢:', False), ('肱動脈最常見', True)])

H2('其他原因')
H3('動脈內藥物注射(故意或意外):血管痙攣、感染性動脈炎、血栓、假性動脈瘤、內膜炎、mycotic aneurysm。')
H3('惰性顆粒/藥物結晶 → 末端動脈阻塞 → 指端壞疽。')
H3([('長期 vasopressor:可致動脈缺血,', False), ('PAD 病人需注意', True)])
H3([('主動脈剝離', True), ( ' 可延伸至鎖骨下/髂股 → 神經缺損或肢體缺血。', False)])

H1('臨床特徵')
H2('六大 P (6 Ps)')
H3([('Pain', True), ( '(最早)', False)])
H3([('Pallor', True)])
H3([('Paralysis', True)])
H3([('Pulselessness', True)])
H3([('Paresthesias', True)])
H3([('Poikilothermia', True), ( '(冷)', False)])
H3([('缺一個 P 不能排除缺血', True)])

H2('進展')
H3('Pain 為最早,位於阻塞遠端。')
H3('皮膚:蒼白 → 斑駁紫紺 → 瘀點/水疱 → 皮膚脂肪壞死。')
H3('血管阻塞:遠端劇烈持續疼痛 + 皮溫下降。')
H3('感覺/運動異常為缺血性神經病變早期表現。')
H3('Two-point discrimination、震動覺、本體覺常先喪失。')
H3([('長期血管病者 ', False), ('遠端脈搏摸不到不具特異性', True), ( ' ,除非合併皮膚改變。', False)])
H3([('原本強脈搏突然消失 → 高度懷疑 ', False), ('急性栓塞', True)])
H3([('Atheroemboli:指端疼痛 + 紫紺、瘀點、肌肉壓痛(', False), ('blue toe syndrome', True), ( ' );脈搏多保留。', False)])
H3('進展性缺血 → 麻木、癱瘓 → 預示壞疽與不可恢復。')
H3([('輕觸覺保留為組織存活良好指標', True)])
H3([('側支循環決定肢體存活,', False), ('「4–6 h 治療窗」並非絕對', True)])

H2('急性 vs. 慢性')
H3([('ALI 定義:症狀 ', False), ('<2 週', True)])
H3('夜間前足痛而醒、需垂腳 → 嚴重動脈阻塞,需立即會診治療。')
H3([('慢性 PAD:', False), ('間歇性跛行 (claudication)', True), ( ',可進展為靜止痛。', False)])
H3([('Claudication:抽筋/痠/疲倦,', False), ('運動誘發、休息緩解', True), ( ' ,2–5 min 消失,固定距離重現。', False)])
H3([('ALI 之疼痛 ', False), ('休息與重力均無法緩解', True), ( ' ,定位不清。', False)])

H2('依阻塞動脈定位跛行')
H3([('Iliac', True), ( ' :臀、大腿(雙側 → 男性陽萎)。', False)])
H3([('Common femoral', True), ( ' :大腿。', False)])
H3([('Superficial femoral', True), ( ' :小腿上 2/3。', False)])
H3([('Popliteal', True), ( ' :小腿下 1/3。', False)])
H3([('Infrapopliteal (tibial, peroneal)', True), ( ' :足部。', False)])

H2('鑑別診斷')
H3([('模擬 ALI 之疾患:', False), ('系統性休克、Phlegmasia cerulea dolens、急性壓迫性神經病變', True)])
H3([('非 PAD 引起之 ALI:', False), ('動脈外傷、主動脈/動脈夾層、動脈炎合併血栓、HIV 動脈病變、高凝、popliteal entrapment、血管痙攣、Compartment syndrome', True)])
H3('其他模擬:脊椎狹窄、神經根壓迫、慢性 compartment syndrome、靜脈跛行、Baker cyst、足/髖/踝關節炎。')
H3([('鑑別三步驟', True), ( ':(1) 找模擬疾患 → (2) 考慮非粥狀硬化原因 → (3) 區分血栓與栓塞。', False)])

H1('診斷')
H2('理學檢查')
H3([('比較雙側皮膚與脈搏;慢性 PAD:', False), ('皮膚光亮、色素沉澱、毛髮喪失、潰瘍、肌肉萎縮、脈搏微弱', True)])
H3('聽診心臟與腹部:雜音與血管雜音。')
H3('POCUS 偵測 AAA。')
H3('手持 Doppler 評估脈搏處血流;10% 病人僅有一條 pedal/ankle 脈搏。')

H2('ABI (Ankle-Brachial Index)')
H3([('<0.9', True), ( ' :慢性 PAD', False)])
H3([('<0.4', True), ( ' :嚴重', False)])
H3([('>1.3', True), ( ' :血管不可壓縮(常為糖尿病鈣化)', False)])

H2('Rutherford 分期(必背)')
H3([('Class I — Viable', True), ( ' :無立即威脅;無感覺/肌力缺損;動靜脈 Doppler 均可聽。', False)])
H3([('Class IIa — Marginally threatened', True), ( ' :及時治療可挽救;微感覺缺損(腳趾);無肌力減弱;動脈 Doppler 不可聽。', False)])
H3([('Class IIb — Immediately threatened', True), ( ' :立即再血管化方可挽救;>腳趾感覺缺損 + 靜止痛;輕至中度肌力減弱;動脈 Doppler 不可聽。', False)])
H3([('Class III — Irreversible', True), ( ' :不可回復;深度麻木、癱瘓僵直;動靜脈均不可聽。', False)])

H2('栓塞 vs. 血栓鑑別')
H3('跛行史:栓塞無、血栓有。')
H3('對側肢體:栓塞正常、血栓有雙側病變。')
H3('來源:栓塞常可找到 (Afib、MI)、血栓無。')
H3('時機:栓塞突發精準、血栓較漸進。')
H3('影像:栓塞 — 銳利截斷、無側支、血管病變少;血栓 — 廣泛病變、有側支、漸窄。')

H2('實驗室')
H3([('細胞缺血/損傷:', False), ('CK、myoglobin、lactate', True)])
H3('代謝:electrolytes、glucose。')
H3('腎功能:BUN、Cr、U/A。')
H3('CBC、PT/PTT。')
H3('心肌標記與 ECG 找誘因(梗塞、心律不整)。')

H2('影像')
H3('CLI 為時效性診斷;懷疑時即會診外科,不必等影像。')
H3([('Duplex US', True), ( ' 對完整與不完整阻塞高度準確,小腿層級準確度下降。', False)])
H3('POCUS 可定位主動脈、髂、股動脈分叉栓子/血栓。')
H3('TTE 找心源性栓塞;TEE 看主動脈根。')
H3('動脈攝影常於 OR/IR 治療前進行。')
H3([('CT 含對比劑為急診最易取得', True), ( ' ;大血管敏感性與傳統血管攝影相當。', False)])
H3('MRI 敏感與特異性最高,但取得性差。')

H1('治療')
H2('ED 醫療治療')
H3([('目標:', False), ('恢復血流以挽救肢體與生命', True), ( ' ,', False), ('預防再次血栓/栓塞', True)])
H3([('Unfractionated Heparin 為首選', True), ( ':', False), ('80 U/kg bolus → 18 U/kg/h', True), ( ' infusion。', False)])
H3('預防血栓延伸、再栓塞、靜脈血栓、阻塞遠端微血栓、再灌流後再阻塞。')
H3([('Direct thrombin inhibitor', True), ( ' (lepirudin、argatroban):懷疑 HIT + 血栓時。', False)])
H3([('Aspirin 325 mg', True), ( '(初劑):抗血小板輔助。', False)])
H3('治療其他低流量狀態(如 HF)以優化灌流。')
H3([('保護缺血肢體免於溫度極端', True)])

H2('依 Rutherford 分期處置')
H3([('Stage I 與 IIa', True), ( ' :可先進行影像檢查再治療。', False)])
H3([('Stage IIb', True), ( ':', False), ('立即再血管化,不額外延遲做影像', True)])
H3([('Stage III', True), ( ':不可逆損傷,', False), ('通常需截肢', True)])
H3([('治療方式:', False), ('catheter-directed thrombolysis、percutaneous mechanical thromboembolectomy、阻塞 bypass graft 修復、percutaneous transluminal angioplasty、外科血管重建', True)])
H3('Catheter-directed intra-arterial thrombolysis 比系統性 IV thrombolysis 更常用。')
H3('亞急性血栓且側支良好者:多採內科治療。')

H2('長期治療(非手術)')
H3([('戒菸、結構化運動、藥物', True), ( '三者並行。', False)])
H3([('抗血小板:', False), ('Aspirin 75–100 mg/d 或 Clopidogrel 75 mg/d', True), ( ',降心血管死亡。', False)])
H3([('不建議起始即雙重抗血小板', True)])
H3([('AHA/ACC 2016 指引:', False), ('Cilostazol', True), ( '(PDE inhibitor) 為間歇性跛行 ', False), ('Class I 建議', True)])
H3([('Pentoxifylline 已不再建議', True)])
H3('多專科團隊(專科醫師、傷口照護、營養師)合作降低罹病率。')

H1('處置與後續')
H2('住院/門診')
H3([('ALI 或慢性缺血急性惡化 → ', False), ('觀察、住院或轉院至有血管外科能力之中心', True)])
H3('慢性 PAD 無立即威脅 + 無其他急性疾病 → 出院,門診追蹤血管外科或初級照護。')
H3([('衛教:症狀惡化(尤其疼痛)立即返診;若無禁忌,', False), ('Aspirin 325 mg 起始,後續每日 81 mg', True)])

H1('特殊族群:上肢缺血')
H2('特性')
H3([('上肢動脈急性阻塞 ', False), ('遠少於下肢', True)])
H3([('因肩、肘周邊側支循環發達,缺血耐受性較佳;', False), ('靜止痛或壞疽罕見,除非有遠端栓塞', True)])

H2('原因與處置')
H3('原因:血管痙攣、動脈炎、外傷、粥狀硬化斑塊破裂、栓塞、醫源性損傷(如肱動脈導管)、胸廓出口症候群、動脈瘤、高凝狀態。')
H3('檢查:臨床評估、肘上下節段血壓、Doppler、duplex US、動脈攝影。')
H3([('威脅肢體之手與前臂缺血治療:', False), ('Heparin + 緊急外科 thromboembolectomy', True)])

# Save
doc.save(OUT)
print(f'Saved: {OUT}')
