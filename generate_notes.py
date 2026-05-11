#!/usr/bin/env python3
"""Generate Tintinalli Ch 59-61 Chinese translation Word document."""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_chinese_font(run, font_name='Microsoft JhengHei', size=11):
    run.font.name = font_name
    run.font.size = Pt(size)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        size = 18
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    elif level == 2:
        size = 14
        run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    elif level == 3:
        size = 12
        run.font.color.rgb = RGBColor(0x44, 0x54, 0x6A)
    else:
        size = 11
    set_chinese_font(run, size=size)
    return p


def add_bullet(doc, segments, level=0):
    """segments: list of (text, bold) tuples"""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.75 + level * 0.75)
    for text, bold in segments:
        run = p.add_run(text)
        run.bold = bold
        set_chinese_font(run, size=11)
    return p


def add_para(doc, segments, indent=0):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    for text, bold in segments:
        run = p.add_run(text)
        run.bold = bold
        set_chinese_font(run, size=11)
    return p


doc = Document()

# Set default style
style = doc.styles['Normal']
style.font.name = 'Microsoft JhengHei'
style.font.size = Pt(11)

# ============================================================
# Title
# ============================================================
title = doc.add_paragraph()
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run = title.add_run('Tintinalli 急診醫學  第 59–61 章重點筆記')
run.bold = True
set_chinese_font(run, size=20)
run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

sub = doc.add_paragraph()
sub.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
r = sub.add_run('（中文翻譯整理版 ‧ 適用住院醫師臨床參考）')
set_chinese_font(r, size=11)

doc.add_paragraph()

# ============================================================
# CHAPTER 59
# ============================================================
add_heading(doc, '第 59 章  主動脈剝離及相關主動脈症候群 (Aortic Dissection and Related Aortic Syndromes)', level=1)

add_heading(doc, '一、前言與流行病學 (Introduction and Epidemiology)', level=2)
items = [
    [('急性主動脈症候群涵蓋多種致命性主動脈急症,包括 ', False), ('主動脈剝離 (aortic dissection)', True), (' 、', False), ('穿透性動脈粥狀硬化潰瘍 (penetrating atherosclerotic ulcer, PAU)', True), (' 、', False), ('壁內血腫 (intramural hematoma, IMH)', True), (' ,以及主動脈瘤滲漏及破裂性腹主動脈瘤(見第 60 章)。', False)],
    [('急性主動脈症候群不常見,但通常致命。發生率約為每年 ', False), ('每 10 萬人 2.9 至 4.7 例', True), ('。', False)],
    [('手術病人的 1 年、5 年、10 年存活率分別為 ', False), ('92%、77%、57%', True), ('。', False)],
    [('22% 的個案在死亡前未被診斷', True)],
    [('Marfan 症候群最常見的心血管併發症為 ', False), ('主動脈根部疾病與 type A 剝離', True), (' (升主動脈)。', False)],
    [('辨識相關基因突變(如 ', False), ('TGFBR2 與 FBN1', True), (' )並合併定期追蹤,可降低致命性結果。', False)],
]
for it in items:
    add_bullet(doc, it)

add_heading(doc, '二、病理生理學 (Pathophysiology)', level=2)
items = [
    [('急性主動脈症候群發生於 ', False), ('慢性高血壓', True), (' 以及其他導致主動脈中膜退化的因素背景下。', False)],
    [('易感因子:', False), ('二葉式主動脈瓣 (bicuspid aortic valve)、Marfan 症候群、Ehlers-Danlos 症候群、家族性主動脈剝離史', True), ('。', False)],
    [('慢性使用 ', False), ('cocaine 或 amphetamine', True), (' 會加速動脈粥狀硬化並提升剝離風險。', False)],
    [('既往 ', False), ('心臟手術', True), (' 亦為主動脈剝離之危險因子。', False)],
    [('所有機轉皆牽涉中膜層的削弱及內膜壁應力增加。', False)],
    [('主動脈剝離', True), (' 發生於內膜撕裂後,血液進入中膜並於內膜與外膜之間剝離。', False)],
    [('兩個最常見的內膜撕裂部位:', False)],
]
for it in items:
    add_bullet(doc, it)
add_bullet(doc, [('升主動脈起始之 ', False), ('sinotubular junction (50–65%)', True)], level=1)
add_bullet(doc, [('左鎖骨下動脈遠端 ', False), ('(20–30%)', True), (' ,即升主動脈與降主動脈交界處。', False)], level=1)

items = [
    [('剝離的血柱形成 ', False), ('false lumen', True), (' ,可向遠端(最常見)、近端、或雙向延伸。', False)],
    [('血液若再次穿入內膜,可重建血流,並可能 ', False), ('錯誤地暗示症狀自發性緩解', True), ('。', False)],
    [('若剝離穿入外膜,則 ', False), ('幾乎都迅速致命', True), ('。', False)],
    [('主動脈剝離有 ', False), ('雙峰年齡分布', True), (' :年輕族群多伴隨結締組織疾病;>50 歲族群多伴隨慢性高血壓及缺血性心臟病。', False)],
    [('既往主動脈剝離為再次剝離的危險因子。', False)],
    [('分類系統:', False), ('Stanford 與 DeBakey 兩種', True), ('。', False)],
    [('Stanford A', True), (' :任何牽涉到升主動脈的剝離。', False)],
    [('Stanford B', True), (' :僅限降主動脈。', False)],
    [('DeBakey 1', True), (' :同時包含升主動脈、主動脈弓、降主動脈。', False)],
    [('DeBakey 2', True), (' :僅升主動脈。', False)],
    [('DeBakey 3', True), (' :僅降主動脈。', False)],
    [('壁內血腫 (IMH)', True), (' 起源於主動脈中膜梗塞,通常因 vasa vasorum 受損所致;可自行緩解或進展為剝離。', False)],
    [('穿透性動脈粥狀硬化潰瘍', True), (' 可導致 IMH、剝離、或主動脈穿孔。', False)],
    [('IMH 約 ', False), ('62% 發生於女性', True), (' ;女性剝離較好發於高齡,合併院內併發症與較高的院內死亡率。', False)],
]
for it in items:
    add_bullet(doc, it)

add_heading(doc, '三、臨床特徵 (Clinical Features)', level=2)
add_heading(doc, '【病史】', level=3)
items = [
    [('初始內膜破裂部位常決定起始症狀;隨剝離延伸或涉及其他動脈/器官,症狀可能改變。', False)],
    [('典型表現:', False), ('突發', True), (' 、', False), ('劇烈的胸痛', True), (' ,放射至 ', False), ('兩肩胛之間', True), (' ,並伴有瀕死感。', False)],
    [('464 例剝離的研究:60% 有前胸痛(Stanford A 較多);腹痛較多見於 Stanford B。', False)],
    [('大多數病人形容疼痛為 ', False), ('一生中最劇烈', True), (' ,', False), ('64% 為銳痛 (sharp)', True), (' ,', False), ('50% 為撕裂 (tearing) 或撕扯感 (ripping)', True), ('。', False)],
    [('亞洲一份研究顯示 ', False), ('高達 40% 病人無胸痛', True), ('。', False)],
    [('Syncope 發生率約 10%', True), (' ,在 Stanford A 較常見。', False)],
    [('22% 剝離發生於既往心臟手術病人。', False)],
    [('靠近頸動脈的剝離可呈現典型中風;', False), ('20% type A 病人有神經學徵候', True), (' ,預後較差。', False)],
    [('脊髓血流中斷可導致 ', False), ('截癱', True), ('。', False)],
    [('進一步遠端剝離可表現為 ', False), ('背痛、側腹痛或腹痛', True), ('。', False)],
    [('近端剝離至主動脈根部可導致 ', False), ('心包填塞', True), (' ,通常致命。', False)],
]
for it in items:
    add_bullet(doc, it)

add_heading(doc, '【理學檢查】', level=3)
items = [
    [('大多數剝離病人理學檢查 ', False), ('相對正常', True), ('。', False)],
    [('Aortic insufficiency 雜音', True), (' 可出現(32%)。', False)],
    [('橈動脈或股動脈 ', False), ('脈搏不對稱', True), (' (15%)。', False)],
    [('兩臂血壓差 ', False), ('>20 mmHg', True), (' 與剝離獨立相關,但 19% 無剝離的急診病人亦可見此差異。', False)],
    [('高血壓常見 (49%) ,', False), ('低血壓 18–25%', True), (' ,後者預後較差。', False)],
    [('主動脈瘤狀擴張可壓迫鄰近結構(食道、喉返神經、上頸交感神經節),造成 ', False), ('吞嚥困難、聲音沙啞、Horner 症候群', True), ('。', False)],
    [('根據 IRAD,三大臨床類別(背景條件、疼痛、理學檢查)的 12 項特徵與急性剝離相關,構成 ', False), ('Aortic Dissection Detection Risk Score (ADD-RS)', True), (' ,每類有一項計 1 分,總分 0–3。', False)],
]
for it in items:
    add_bullet(doc, it)

add_heading(doc, '【ADD-RS 三大類別】', level=3)
add_para(doc, [('類別 1:背景條件', True)])
for x in ['Marfan 症候群', '主動脈疾病家族史', '主動脈瓣疾病', '近期主動脈操作', '胸主動脈瘤']:
    add_bullet(doc, [(x, False)], level=1)
add_para(doc, [('類別 2:胸、背或腹部疼痛特徵', True)])
for x in ['突發 (abrupt onset)', '強度劇烈 (severe in intensity)', '撕裂或撕扯感 (ripping/tearing)']:
    add_bullet(doc, [(x, False)], level=1)
add_para(doc, [('類別 3:理學檢查異常', True)])
for x in ['四肢脈搏振幅差或 SBP 差', 'Focal 神經缺損 + 胸/背/腹痛', '新發 aortic regurgitation 雜音 + 痛', '休克或低血壓']:
    add_bullet(doc, [(x, False)], level=1)

add_heading(doc, '四、診斷 (Diagnosis)', level=2)
items = [
    [('胸痛鑑別診斷廣泛,加上器官缺血表現多樣,使診斷極具挑戰。', False)],
    [('重要鑑別:', False), ('急性冠心症、心包疾病、中風、肌肉骨骼疾病、脊髓損傷、腹腔內疾病、肺栓塞/肺炎/胸膜炎/氣胸', True), ('。', False)],
    [('缺血表現可能隨時間變化,易誤導醫師。', False)],
    [('誤診相關因素:', False), ('步入式就診、CXR 縱膈寬度/主動脈輪廓正常、雙側脈搏對稱、非特異性症狀', True), ('。', False)],
]
for it in items:
    add_bullet(doc, it)

add_heading(doc, '【ECG】', level=3)
items = [
    [('剝離與 ACS 在 ECG 上不易區分,且剝離可阻塞冠狀動脈血流。', False)],
    [('新 Q 波或 ST elevation:', False), ('3–4%', True), ('。', False)],
    [('ST 下降:15–22%。', False)],
    [('非特異性 ST/T 變化:41–62%。', False)],
    [('ECG ', False), ('完全正常僅 19–31%', True), ('。', False)],
]
for it in items:
    add_bullet(doc, it)

add_heading(doc, '【生物標記】', level=3)
items = [
    [('D-dimer 是最廣泛研究的標記。', False)],
    [('Meta-analysis (cut-off 500 ng/mL):敏感度 ', False), ('98%', True), (' ,LR− ', False), ('0.05', True), (' ;特異度低 ', False), ('41%', True), ('。', False)],
    [('指引', False), ('不建議單獨以 D-dimer 排除剝離', True), (' ,false negative 高達 ', False), ('18%', True), (' ,且可能與血小板數高相關。', False)],
    [('ADD-RS 0–1 + D-dimer 陰性(<500 ng/mL):剝離率僅 ', False), ('0.3%', True), (' ,但此研究仍待外部驗證。', False)],
]
for it in items:
    add_bullet(doc, it)

add_heading(doc, '【影像】', level=3)
items = [
    [('CXR', True), (' :', False), ('12–37% 完全正常', True), (' ,故不能用以排除剝離。', False)],
    [('最常見異常:', False), ('縱膈擴大或主動脈輪廓異常', True), (' ;其他:胸水、內膜鈣化移位、氣管/支氣管/食道偏移。', False)],
    [('CT 為診斷影像首選', True), (' ,可清楚顯示 false lumen、剝離瓣、延伸至大血管、破裂徵兆與末端器官損傷。CT 包含含對比與無對比兩階段。', False)],
    [('TEE', True), (' :經驗操作者敏感性與 CT 相當;通常需鎮靜;食道疾病為相對禁忌;氣管/左支氣管之空氣可能干擾升主動脈評估。', False)],
    [('MRI', True), (' :用於穩定病人評估。', False)],
    [('Triple rule-out (CTA 同時評估 CAD、PE、剝離)', True), (' :目前證據不支持其改善診斷率或臨床結果,', False), ('不推薦', True), ('。', False)],
]
for it in items:
    add_bullet(doc, it)

add_heading(doc, '五、治療 (Treatment)', level=2)
add_heading(doc, '【降血壓藥:負性肌力藥物優先】', level=3)
items = [
    [('剝離可導致低血壓,需液體/血品復甦;但多數情況下需要降壓治療。', False)],
    [('初始治療應為 ', False), ('負性肌力藥物', True), (' ,以降低血壓而不增加主動脈壁的剪力。', False)],
    [('β-blocker 為首選', True), (' ,短效如 ', False), ('esmolol 或 labetalol', True), (' 優於長效。', False)],
    [('理想目標 SBP 約 ', False), ('120–130 mmHg', True), (' ;部分指引建議 100–120 mmHg(無 AR 情況下)。', False)],
    [('Esmolol', True), (' :起始 0.1–0.5 mg/kg IV 推注 1 分鐘,接續 0.025–0.2 mg/kg/min 輸注。', False)],
    [('Labetalol', True), (' :起始 10–20 mg IV,每 10 分鐘重複 20–40 mg,直至目標或最大 300 mg。', False)],
    [('β-blocker 使用與 IRAD 資料中改善存活率相關', True), ('。', False)],
]
for it in items:
    add_bullet(doc, it)

add_heading(doc, '【血管擴張劑】', level=3)
add_bullet(doc, [('在成功使用 esmolol 或 labetalol 之後可加入血管擴張劑。', False)])
add_bullet(doc, [('選項:', False), ('IV nicardipine、clevidipine、nitroglycerin、nitroprusside', True), ('(見第 57、58 章)。', False)])

add_heading(doc, '六、處置與後續 (Disposition and Follow-up)', level=2)
items = [
    [('儘速轉介給 ', False), ('血管外科', True), ('。', False)],
    [('多數病人仍以開放性修復為主流;', False), ('endovascular 修復使用日增', True), ('。', False)],
    [('需 ', False), ('加護病房', True), ('進行血流動力學治療及監測。', False)],
    [('急性 IMH 與穿透性潰瘍的自然病程仍不清楚,治療仍有爭議。', False)],
    [('任何急性主動脈症候群病人 ', False), ('不應未經專科會診即出院', True), ('。', False)],
]
for it in items:
    add_bullet(doc, it)

add_heading(doc, '七、特殊族群:懷孕合併剝離', level=2)
items = [
    [('懷孕期主動脈剝離罕見,', False), ('多發生於第三孕期及產後', True), ('。', False)],
    [('危險因子:', False), ('二葉式主動脈瓣、結締組織疾病、高血壓、家族史', True), ('。', False)],
    [('Marfan 症候群孕婦中,', False), ('4.4% 妊娠合併剝離', True), ('。', False)],
    [('診斷一旦考量,需 ', False), ('同時會診產科及心血管外科', True), ('。', False)],
]
for it in items:
    add_bullet(doc, it)

doc.add_page_break()

# ============================================================
# CHAPTER 60
# ============================================================
add_heading(doc, '第 60 章  動脈瘤疾病 (Aneurysmal Disease)', level=1)

add_heading(doc, '一、前言 (Introduction)', level=2)
items = [
    [('動脈瘤', True), (' :動脈壁擴張至正常直徑的 ', False), ('1.5 倍以上', True), ('。', False)],
    [('經典分類:', False), ('true aneurysm、pseudoaneurysm、mycotic aneurysm', True), ('。', False)],
    [('True aneurysm', True), (' 之動脈壁包含血管所有層次。', False)],
    [('危險因子:', False), ('抽菸、年齡增加、白種人、高血壓、高血脂、結締組織疾病、動脈瘤家族史', True), ('。', False)],
]
for it in items:
    add_bullet(doc, it)

add_heading(doc, '二、病理生理 (Pathophysiology)', level=2)
items = [
    [('Elastin、collagen 及 fibrolamellar units 漸減,中膜變薄,張力下降。', False)],
    [('動脈擴張與壁張力相互強化 ', False), ('(Laplace law:wall tension = pressure × radius)', True), (' ,促使更大擴張。', False)],
    [('動脈瘤擴張率多變,大動脈瘤擴張更快(平均 ', False), ('每年 0.25–0.5 cm', True), (' ),但突發擴張無法預測。', False)],
    [('破裂風險隨體積增加而上升;', False), ('破裂為災難性事件', True), ('。', False)],
    [('社經地位較低者易以破裂表現,且手術後預後較差。', False)],
    [('Pseudoaneurysm', True), (' 之壁部分為血管壁、部分為纖維或周邊組織;可見於導管插入處、外科吻合處、外傷後、感染。', False)],
    [('Mycotic aneurysm', True), (' :續發於瓣膜內膜炎的敗血性栓塞,常見於免疫低下者。', False)],
    [('Infected aneurysm', True), (' :既存動脈瘤發生菌血症或直接感染擴散。', False)],
    [('IVDU 為兩者之重要危險因子', True), (' ,需注意誤診為蜂窩組織炎或膿瘍。', False)],
    [('周邊與內臟動脈瘤較少見,但臨床重要。', False)],
    [('膕動脈瘤 (popliteal artery aneurysm) 為最常見之周邊動脈瘤', True), (' ,常合併對側膕動脈或腹主動脈瘤。', False)],
    [('內臟動脈瘤多為靜默,常因破裂等併發症被發現。', False)],
    [('併發症:', False), ('破裂(致死率 80%)', True), (' ,', False), ('血栓造成器官缺血', True), ('。', False)],
]
for it in items:
    add_bullet(doc, it)

add_heading(doc, '三、動脈瘤的一般臨床特徵', level=2)
items = [
    [('臨床表現多 ', False), ('非特異性', True), (' ,常依位置、對鄰近結構的壓迫或周邊栓塞而異。', False)],
    [('內臟動脈瘤常於腹/側腹影像檢查中偶發發現;下肢者常於 Doppler 檢查時發現。', False)],
    [('軀幹動脈瘤一旦破裂,', False), ('迅速進入出血性休克', True), (' ,若無立即手術,死亡率高。', False)],
]
for it in items:
    add_bullet(doc, it)

add_heading(doc, '四、有症狀腹主動脈瘤 (Symptomatic AAA)', level=2)
items = [
    [('AAA 定義:腹主動脈 ', False), ('≥3.0 cm', True), (' ;手術修復考量 ', False), ('≥5.0 cm', True), ('。', False)],
    [('一等親有 AAA 者佔 18%。', False)],
    [('多數 ', False), ('>60 歲', True), (' 及男性風險增加。', False)],
    [('抽菸為最重要之環境危險因子', True), (' ,長期抽菸者 AAA 盛行率 ', False), ('超過終生不吸菸者四倍', True), ('。', False)],
]
for it in items:
    add_bullet(doc, it)

add_heading(doc, '【臨床特徵】', level=3)
items = [
    [('最常見症狀:', False), ('背痛或腹痛', True), (' ,突發、劇烈;約半數為撕裂或撕扯感。', False)],
    [('經典三聯徵', True), (' :', False), ('腹痛 + 搏動性腹部腫塊 + 低血壓', True), (' ,', False), ('僅 1/3 病人具備', True), ('。', False)],
    [('Syncope 約 ', False), ('10%', True), ('。', False)],
    [('非典型疼痛部位:側腹、鼠蹊、單側腹部象限、髖部。', False)],
    [('其他症狀:噁心嘔吐、膀胱痛、髖痛、裡急後重。', False)],
    [('11% 病人否認疼痛', True), ('。', False)],
    [('可能表現:消化道出血(', False), ('aortoenteric fistula', True), (' )、肢端缺血(血栓栓塞)、休克、猝死。', False)],
    [('猝死多為腹腔內破裂導致大量出血。', False)],
    [('Syncope 無預警 + 隨後劇烈腹/背痛 → 考慮 ', False), ('破裂後暫時被局限', True), ('。', False)],
]
for it in items:
    add_bullet(doc, it)

add_heading(doc, '【診斷】', level=3)
items = [
    [('臨床懷疑後,若情況允許再以影像確認。', False)],
    [('觸診敏感性依直徑而異:3.0–3.9 cm = ', False), ('29%', True), (' ;4.0–4.9 cm = ', False), ('50%', True), (' ;≥5.0 cm = ', False), ('76%', True), ('。', False)],
    [('觸壓痛常被解讀為擴大或破裂的徵兆,但 ', False), ('無壓痛不代表主動脈完整', True), ('。', False)],
    [('肥胖者觸診困難;消瘦者主動脈尺寸可能被高估。', False)],
    [('低血壓與心搏過速僅少數病人具備。', False)],
    [('鑑別診斷:暈厥、腹痛、胸痛、背痛、休克之原因;包括心臟、腹腔、後腹腔疾病。', False)],
    [('當突發背痛 + 暈厥或休克時必須考慮主動脈瘤破裂。', True)],
    [('小心 ', False), ('誤診為腎結石', True), ('。', False)],
    [('合併疾病(冠心病、慢性肺病)可能轉移醫師注意。', False)],
    [('Cullen sign', True), (' (臍周瘀斑)、', False), ('Grey Turner sign', True), (' (側腹瘀斑)罕見。', False)],
    [('後腹腔出血可至會陰/鼠蹊,造成陰囊或外陰血腫、鼠蹊腫塊。', False)],
    [('Iliopsoas sign:伸髖時疼痛。', False)],
    [('AAA 通常 ', False), ('不影響股動脈搏動', True), ('。', False)],
    [('Aortoenteric fistula', True), (' :無肝病的不明原因/大量上下消化道出血;主動脈移植史增加風險;多侵犯十二指腸。', False)],
    [('Aortovenous fistula:可造成高輸出心衰竭、遠端動脈血流降低、中央靜脈容積上升。', False)],
    [('Contained chronic rupture 少見,後腹腔纖維化可限制出血;發炎反應使疼痛持續。', False)],
]
for it in items:
    add_bullet(doc, it)

add_heading(doc, '【影像】', level=3)
items = [
    [('不穩定病人不可送出急診做影像', True), ('。', False)],
    [('床邊 US 為不穩定病人首選', True), ('。', False)],
    [('平片有時偶然發現動脈瘤;65% 症狀性 AAA 主動脈鈣化,側位較明顯;', False), ('平片不能排除 AAA', True), ('。', False)],
    [('Bedside US:敏感性 ', False), ('>90%', True), (' ;肥胖、腸氣、腹痛影響操作。', False)],
    [('測量:', False), ('外壁到外壁', True), (' ,橫切與縱切兩平面;辨識 SMA 以區分主動脈與下腔靜脈。', False)],
    [('主動脈直徑 ', False), ('<3.0 cm 可排除急性動脈瘤疾病', True), ('。', False)],
    [('CT 含對比劑為穩定病人之首選詳細影像', True), ('。', False)],
    [('無法使用對比劑者:無強化 CT 仍可顯示動脈瘤大小與後腹腔血腫。', False)],
]
for it in items:
    add_bullet(doc, it)

add_heading(doc, '【治療】', level=3)
items = [
    [('關鍵:辨識動脈瘤並 ', False), ('儘速安排手術控制出血', True), ('。', False)],
    [('臨床懷疑時立即會診外科,', False), ('低血壓或末端灌注不足時不可延遲', True), (' ;床邊 US 即可初步診斷。', False)],
    [('進階影像僅用於穩定且不太可能為破裂之病人。', False)],
    [('標準急救:兩條大號 IV、心電監測、補氧。', False)],
    [('積極輸液可能有害', True), (' ;血管外科學會建議 ', False), ('允許性低血壓 (permissive hypotension)', True), (' :SBP 目標 ', False), ('80–90 mmHg', True), (' ,意識清楚為輸液目標。', False)],
    [('Esmolol(半衰期 9 分鐘):用於懷疑擴大且嚴重高血壓時,可滴定至 ', False), ('SBP 120 mmHg', True), (' ;血壓突降可迅速停用。', False)],
    [('REBOA', True), (' (主動脈內球囊復甦):用於 AAA 破裂阻斷上方主動脈;現有研究小且死亡率改善不明;需訓練及術後管理計畫;目前主要用於外傷情境。', False)],
    [('任何大小之有症狀動脈瘤皆視為緊急', True), (' ,若無手術能力應啟動復甦並轉院。', False)],
    [('立即轉院並由能辨識/治療休克者陪伴', True), (' (備有液體、血品)。', False)],
    [('有症狀者圍術期死亡風險為無症狀者的 ', False), ('2 倍', True), (' ,破裂者達 ', False), ('7 倍', True), ('。', False)],
    [('所有無症狀者皆需安排追蹤。', False)],
    [('≥5 cm 風險上升 → 數日內追蹤', True), (' ,並衛教病人惡化或暈厥立即返診。', False)],
    [('女性 AAA 較小尺寸即可能破裂', True), (' ;懷孕及產後破裂風險上升。', False)],
    [('3–5 cm 由初級照護醫師或外科追蹤,血壓控制。', False)],
    [('男性 <5.5 cm 修復未改善存活,女性可能適合。', False)],
    [('EVAR', True), (' 為較不侵入之選項;', False), ('併發症需注意', True), (' :graft leak、阻塞、腎功能不全、aortoenteric fistula、腸缺血。', False)],
]
for it in items:
    add_bullet(doc, it)

add_heading(doc, '五、胸主動脈瘤 (Thoracic Aortic Aneurysm)', level=2)
items = [
    [('可壓迫或侵蝕鄰近結構(食道、氣管、支氣管、神經)。', False)],
    [('一旦侵蝕鄰近結構,', False), ('通常立即致死', True), (' ,除非有迅速復甦。', False)],
    [('無症狀者應轉介外科及初級照護,', False), ('血壓控制', True), ('。', False)],
    [('修復可採開放性或 endovascular ,兩者長期結果仍具爭議。', False)],
]
for it in items:
    add_bullet(doc, it)

add_heading(doc, '六、肢體與內臟動脈瘤', level=2)
items = [
    [('肢體動脈瘤(膕、鎖骨下、股動脈)可表現為肢體威脅性併發症,包括 ', False), ('破裂、血栓、周邊栓塞', True), ('。', False)],
    [('內臟動脈瘤通常靜默至併發症發生;包括 ', False), ('腎血管性高血壓、腎動脈血栓、遠端栓塞造成器官梗塞、動靜脈瘻、破裂', True), ('。', False)],
]
for it in items:
    add_bullet(doc, it)

add_heading(doc, '【膕動脈瘤 (Popliteal aneurysm)】', level=3)
for it in [
    [('定義:膕動脈局部擴大 >2 cm 或 >正常 150%。', False)],
    [('症狀:膕窩不適、合併 DVT 之腿腫、間歇性跛行。', False)],
    [('破裂罕見', True), (' ;最嚴重併發症為血栓/栓塞造成 ', False), ('急性肢體缺血', True), ('。', False)],
    [('觸診膕窩可見硬塊或搏動性腫塊;', False), ('US 為診斷工具', True), (' ,需評估雙側。', False)],
]:
    add_bullet(doc, it)

add_heading(doc, '【股動脈/髂動脈瘤】', level=3)
for it in [
    [('表現:鼠蹊或大腿上端搏動性腫塊、陰囊血腫、急性肢體缺血。', False)],
    [('髂動脈瘤難以臨床懷疑,因髂動脈不能直接觸診,症狀可像泌尿、腸道或鼠蹊疾患。', False)],
]:
    add_bullet(doc, it)

add_heading(doc, '【肝動脈瘤】', level=3)
for it in [
    [('原因:動脈粥狀硬化或外傷。', False)],
    [('破裂可表現為急性出血性休克(腹腔或後腹腔)。', False)],
    [('Quincke 三聯徵', True), (' :', False), ('黃疸、膽絞痛、上消化道出血', True), (' (hemobilia 由肝動脈瘤滲漏至膽管造成)。', False)],
]:
    add_bullet(doc, it)

add_heading(doc, '【脾動脈瘤】', level=3)
for it in [
    [('表現:左上腹痛、未分化休克、腹腔內出血。', False)],
    [('破裂預後差(腹腔位置 + 表現非特異)。', False)],
    [('特別注意 ', False), ('懷孕第三孕期', True), ('。', False)],
    [('門脈高壓亦為危險因子。', False)],
]:
    add_bullet(doc, it)

add_heading(doc, '【鎖骨下與無名動脈瘤】', level=3)
add_bullet(doc, [('可致上肢缺血:上肢疼痛、無力、蒼白、感覺異常,需考慮此診斷。', False)])

add_heading(doc, '【吻合動脈瘤】', level=3)
for it in [
    [('可發生於主動脈、髂、股動脈;破裂可致大出血,可有 ', False), ('警示性小出血', True), ('(疼痛或短暫低血壓)。', False)],
    [('可侵蝕鄰近腸道 → ', False), ('aortoenteric fistula', True), (' ;AAA 修復史 + 嚴重下消化道出血為經典表現。', False)],
]:
    add_bullet(doc, it)

add_heading(doc, '【影像與治療】', level=3)
for it in [
    [('Doppler US 為周邊動脈瘤(股、膕、鎖骨下)之初步無創檢查。', False)],
    [('合併缺血或血栓 → 緊急 ', False), ('CT angiography', True), (' 兼具診斷與治療規劃。', False)],
    [('有症狀者需迅速診斷與會診,可能威脅肢體。', False)],
    [('臨床有缺血徵象者即刻會診血管外科。', False)],
    [('無症狀周邊動脈瘤由血管外科門診處置。', False)],
    [('Mycotic/infected aneurysm', True), (' :需 ', False), ('長期 IV 抗生素 + 手術', True), (' ;懷疑破裂者立即會診。', False)],
]:
    add_bullet(doc, it)

doc.add_page_break()

# ============================================================
# CHAPTER 61
# ============================================================
add_heading(doc, '第 61 章  動脈阻塞 (Arterial Occlusion)', level=1)

add_heading(doc, '一、前言與流行病學', level=2)
items = [
    [('肢體缺血', True), (' :血流不足以致組織灌注降低、產生症狀。', False)],
    [('急性肢體缺血 (acute limb ischemia, ALI)', True), (' :新發或進展之血流缺損,需立即辨識與治療以挽救肢體。', False)],
    [('臨界肢體缺血 (critical limb ischemia, CLI)', True), (' :周邊動脈疾病之末期,出現靜止痛、潰瘍或壞疽。', False)],
    [('最重要危險因子:', False), ('抽菸與糖尿病', True), ('。', False)],
    [('其他:高血脂、高血壓、高同半胱胺酸、C-reactive protein 上升。', False)],
    [('周邊動脈疾病嚴重度與心肌梗塞、缺血性中風、血管性死亡風險相關。', False)],
    [('最常受影響動脈:', False), ('股膕、脛、主動脈髂、頭臂動脈', True), ('(依次)。', False)],
    [('CLI 1 年死亡率達 ', False), ('25%', True), (' ,', False), ('1/4 存活者需截肢', True), ('。', False)],
    [('全球 PAD 患者超過 2 億;美國 CLI 年發生率約 0.35% ,平均盛行率 1.33%。', False)],
]
for it in items:
    add_bullet(doc, it)

add_heading(doc, '二、病理生理', level=2)
items = [
    [('ALI 為動脈供血急遽下降,組織灌注不足而威脅肢體存活。', False)],
    [('無側支循環時,周邊神經與骨骼肌可於 ', False), ('4–6 小時', True), (' 內發生不可逆變化。', False)],
    [('血流恢復後可出現 ', False), ('再灌流損傷', True), (' :compartment syndrome、橫紋肌溶解、代謝紊亂、高血鉀、肌紅蛋白血症、代謝性酸中毒、CK 上升。', False)],
    [('再灌流的損傷與阻塞時間、位置、側支循環、原本肢體健康狀況相關。', False)],
    [('1/3 動脈阻塞死亡源於再血管化後代謝併發症', True), ('。', False)],
]
for it in items:
    add_bullet(doc, it)

add_heading(doc, '【血栓 (Thrombosis)】', level=3)
items = [
    [('最常見原因', True), (' :可發生於原生血管或 bypass graft。', False)],
    [('下肢:>80% 之 ALI 為血栓性。', False)],
    [('上肢:約一半為血栓性,1/3 為栓塞。', False)],
    [('臨床上不易區分血栓與栓塞。', False)],
    [('血栓常見於 ', False), ('動脈粥狀硬化', True), (' 背景;因有側支循環,缺血表現較漸進。', False)],
    [('斑塊破裂亦可直接栓塞或誘發栓塞性血栓。', False)],
    [('進展機轉:(1) 血栓延伸阻塞側支;(2) 缺血性水腫造成 compartment syndrome;(3) 血栓碎片進入微循環;(4) 微血管細胞水腫。', False)],
    [('大血管再灌流不一定能解除微循環阻塞。', True)],
    [('正常血管中發生動脈血栓 → 需評估高凝狀態。', False)],
]
for it in items:
    add_bullet(doc, it)

add_heading(doc, '【栓塞 (Embolism)】', level=3)
items = [
    [('心臟為周邊栓塞之主要來源,', False), ('心房顫動 (Afib) 最常見', True), ('。', False)],
    [('另一來源:近期 MI 之 ventricular mural thrombus(24h 至 3 個月之間,整體 3%,前壁 MI 高達 9%)。', False)],
    [('Afib 與 MI 都會造成心壁活動降低與血流滯留,促進血栓形成。', False)],
    [('其他來源:機械瓣膜、心房黏液瘤、瓣膜贅生物、假體心臟裝置碎片。', False)],
    [('非心源性:動脈瘤與動脈粥樣斑塊之血栓(主動脈髂、股、膕、鎖骨下)。', False)],
    [('Atheroembolism', True), (' :膽固醇與血小板聚合物造成微循環阻塞 → ', False), ('blue toe syndrome', True), (' 、手部表現、TIA。', False)],
    [('Paradoxical embolization', True), (' :靜脈血栓經心內分流(常為卵圓孔未閉)至左側循環。', False)],
    [('栓子常停滯於血管 ', False), ('狹窄或分叉處', True), (' ;下肢最常見:', False), ('股總動脈分叉處 > 膕動脈', True), ('。', False)],
    [('上肢:', False), ('肱動脈最常見', True), ('。', False)],
]
for it in items:
    add_bullet(doc, it)

add_heading(doc, '【其他原因】', level=3)
items = [
    [('動脈內藥物注射(故意或意外):血管痙攣、感染性動脈炎、血栓、假性動脈瘤、內膜炎、mycotic aneurysm。', False)],
    [('惰性顆粒/藥物結晶可栓塞末端動脈 → 指端壞疽。', False)],
    [('長期使用 vasopressor 可致動脈缺血,', False), ('PAD 患者更需注意', True), ('。', False)],
    [('主動脈剝離', True), ('可延伸至鎖骨下或髂股系統,表現神經缺損或肢體缺血。', False)],
]
for it in items:
    add_bullet(doc, it)

add_heading(doc, '三、臨床特徵 (Clinical Features)', level=2)
items = [
    [('六大 P (Six Ps)', True), (' :', False), ('Pain、Pallor、Paralysis、Pulselessness、Paresthesias、Poikilothermia', True), ('(冷)。', False)],
    [('缺一不能排除缺血', True), ('。', False)],
    [('Pain 為最早症狀,位於阻塞遠端。', False)],
    [('皮膚變化:先蒼白 → 斑駁紫紺 → 瘀點/水疱 → 皮膚脂肪壞死。', False)],
    [('血管阻塞:遠端肢體劇烈持續疼痛 + 皮膚溫度下降。', False)],
    [('感覺/運動異常為缺血性神經病變之早期表現。', False)],
    [('Two-point discrimination、震動覺、本體覺常先喪失。', False)],
    [('在長期血管病者,', False), ('遠端脈搏摸不到不具特異性', True), (' ,除非合併皮膚改變。', False)],
    [('原本強的脈搏突然消失 → 高度懷疑 ', False), ('急性栓塞', True), ('。', False)],
    [('Atheroemboli:指端疼痛 + 紫紺、瘀點、肌肉壓痛(', False), ('blue toe syndrome', True), (' );脈搏多保留。', False)],
    [('進展性缺血:麻木、癱瘓 → 預示壞疽與不可恢復。', False)],
    [('輕觸覺保留為組織存活的良好指標', True), ('。', False)],
    [('側支循環決定肢體存活,', False), ('「4–6 小時治療窗」並非絕對', True), ('。', False)],
]
for it in items:
    add_bullet(doc, it)

add_heading(doc, '【急性 vs. 慢性動脈疾病】', level=3)
items = [
    [('ALI 定義:症狀 ', False), ('<2 週', True), ('。', False)],
    [('夜間前足痛而醒、需垂腳 → 嚴重動脈阻塞,需立即會診治療。', False)],
    [('慢性 PAD:', False), ('間歇性跛行 (claudication)', True), (' ,可能進展為靜止痛。', False)],
    [('Claudication 特徵:抽筋/痠/疲倦,', False), ('運動誘發、休息緩解', True), (' ,2–5 分鐘消失,固定距離重現。', False)],
    [('ALI 之疼痛 ', False), ('休息與重力均無法緩解', True), (' ,定位不清,慢性疼痛可顯著加劇。', False)],
]
for it in items:
    add_bullet(doc, it)

add_heading(doc, '【依阻塞動脈定位跛行】', level=3)
table_items = [
    ('Iliac 動脈', '臀、大腿(偶及小腿;雙側 → 男性陽萎)'),
    ('Common femoral', '大腿'),
    ('Superficial femoral', '小腿上 2/3'),
    ('Popliteal', '小腿下 1/3'),
    ('Infrapopliteal (tibial, peroneal)', '足部'),
]
for art, loc in table_items:
    add_bullet(doc, [(art, True), (':', False), (loc, False)])

add_heading(doc, '【鑑別診斷】', level=3)
items = [
    [('模擬 ALI 之疾患:', False), ('系統性休克、Phlegmasia cerulea dolens、急性壓迫性神經病變', True), ('。', False)],
    [('非 PAD 引起之 ALI:', False), ('動脈外傷、主動脈/動脈夾層、動脈炎合併血栓(GCA、Buerger\'s)、HIV 動脈病變、高凝、popliteal adventitial cyst、popliteal entrapment、血管痙攣合併血栓(ergotism)、Compartment syndrome', True), ('。', False)],
    [('其他模擬:脊椎狹窄、神經根壓迫、慢性 compartment syndrome、靜脈跛行、Baker\'s cyst、足/髖/踝關節炎。', False)],
    [('鑑別三步驟', True), (' :(1) 找模擬疾患;(2) 考慮非粥狀硬化原因;(3) 區分血栓與栓塞。', False)],
]
for it in items:
    add_bullet(doc, it)

add_heading(doc, '四、診斷 (Diagnosis)', level=2)
items = [
    [('比較雙側肢體之皮膚與脈搏:慢性 PAD 表徵 = ', False), ('皮膚光亮、色素沉澱、毛髮喪失、潰瘍、肌肉萎縮、脈搏微弱', True), ('。', False)],
    [('聽診心臟與腹部:雜音與血管雜音。', False)],
    [('POCUS 可協助偵測 AAA。', False)],
    [('手持 Doppler 評估脈搏處之血流;10% 病人僅一條 pedal/ankle 脈搏。', False)],
    [('ABI (Ankle-Brachial Index)', True), (' :', False)],
]
for it in items:
    add_bullet(doc, it)
add_bullet(doc, [('<0.9', True), (' :慢性 PAD', False)], level=1)
add_bullet(doc, [('<0.4', True), (' :嚴重', False)], level=1)
add_bullet(doc, [('>1.3', True), (' :血管不可壓縮(常為糖尿病鈣化)', False)], level=1)

add_heading(doc, '【Rutherford 分期】(必背)', level=3)
items = [
    [('Class I — Viable', True), (' :無立即威脅;無感覺/肌力缺損;動脈與靜脈 Doppler 均可聽見。', False)],
    [('Class IIa — Marginally threatened', True), (' :及時治療可挽救;微感覺缺損(腳趾);無肌力減弱;動脈 Doppler 不可聽。', False)],
    [('Class IIb — Immediately threatened', True), (' :立即再血管化方可挽救;感覺缺損 > 腳趾合併靜止痛;輕至中度肌力減弱;動脈 Doppler 不可聽。', False)],
    [('Class III — Irreversible', True), (' :無法回復;深度麻木、癱瘓僵直;動脈與靜脈均不可聽。', False)],
]
for it in items:
    add_bullet(doc, it)

add_heading(doc, '【栓塞 vs. 血栓鑑別】', level=3)
table2 = [
    ('跛行史', '無', '有'),
    ('對側肢體', '正常', '雙側病變徵象'),
    ('來源', '常可找到 (Afib、MI)', '無'),
    ('時機', '突發、可確切時間', '較漸進'),
    ('影像', '突然銳利截斷、無側支、血管病變少', '廣泛病變、有側支、血流漸窄'),
]
add_para(doc, [('因素', True), ('  |  ', False), ('栓塞', True), ('  |  ', False), ('血栓', True)])
for a, b, c in table2:
    add_bullet(doc, [(a, True), (':', False), ('栓塞 — ', False), (b, False), (' / 血栓 — ', False), (c, False)])

add_heading(doc, '【實驗室】', level=3)
items = [
    [('細胞缺血/損傷標記:', False), ('CK、myoglobin、lactate', True), ('。', False)],
    [('代謝狀況:', False), ('electrolytes、glucose', True), ('。', False)],
    [('腎功能:', False), ('BUN、Cr、urinalysis', True), ('。', False)],
    [('CBC、PT/PTT。', False)],
    [('心肌標記與 ECG 可找誘因(梗塞、心律不整)。', False)],
]
for it in items:
    add_bullet(doc, it)

add_heading(doc, '【影像】', level=3)
items = [
    [('CLI 為時效性診斷;懷疑時立即會診外科,不必等影像。', False)],
    [('Duplex US', True), (' 對完整與不完整阻塞高度準確,小腿層級準確度下降。', False)],
    [('POCUS 可定位主動脈、髂、股動脈分叉的栓子/血栓。', False)],
    [('TTE 找心源性栓塞;TEE 可看主動脈根部。', False)],
    [('動脈攝影常於 OR/IR 進行於治療前。', False)],
    [('CT 含對比劑', True), (' 為急診最易取得;大血管敏感性與傳統血管攝影相當。', False)],
    [('MRI 敏感與特異性最高,但取得性差。', False)],
    [('影像選擇需與血管外科/IR 共同決策。', False)],
]
for it in items:
    add_bullet(doc, it)

add_heading(doc, '五、治療', level=2)
items = [
    [('目標:', False), ('恢復血流以挽救肢體與生命', True), (' ,並 ', False), ('預防再次血栓/栓塞', True), ('。', False)],
    [('IV unfractionated heparin', True), (' 為首選起始治療:', False), ('80 U/kg bolus → 18 U/kg/h 輸注', True), ('。', False)],
    [('可預防血栓延伸、再次栓塞、靜脈血栓、阻塞遠端微血栓、再灌流後再阻塞。', False)],
    [('Direct thrombin inhibitors', True), (' (lepirudin、argatroban):用於懷疑 HIT 合併血栓時。', False)],
    [('Aspirin 325 mg', True), ('(初劑)可透過抗血小板作用加強血栓減少。', False)],
    [('治療其他低流量狀態(如心衰竭)以優化灌注。', False)],
    [('保護缺血肢體免於溫度極端', True), ('。', False)],
]
for it in items:
    add_bullet(doc, it)

add_heading(doc, '【依 Rutherford 分期處置】', level=3)
items = [
    [('Stage I 與 IIa', True), (' :可先進行影像檢查再治療。', False)],
    [('Stage IIb', True), (' :', False), ('立即再血管化,不額外延遲做影像', True), ('。', False)],
    [('Stage III', True), (' :不可逆損傷,通常 ', False), ('需截肢', True), ('。', False)],
    [('治療方式:', False), ('catheter-directed thrombolysis、percutaneous mechanical thromboembolectomy、阻塞 bypass graft 修復、percutaneous transluminal angioplasty、外科血管重建', True), ('。', False)],
    [('Catheter-directed intra-arterial thrombolysis 現在比系統性 IV thrombolysis 更常用。', False)],
    [('亞急性血栓且側支良好者:多採內科治療。', False)],
]
for it in items:
    add_bullet(doc, it)

add_heading(doc, '【長期治療(非手術)】', level=3)
items = [
    [('戒菸、結構化運動、藥物', True), ('三者並行。', False)],
    [('抗血小板:', False), ('Aspirin 75–100 mg/d 或 Clopidogrel 75 mg/d', True), (' ,可降低心血管死亡。', False)],
    [('不建議起始即雙重抗血小板。', True)],
    [('AHA/ACC 2016 指引:', False), ('Cilostazol', True), ('(磷酸二酯酶抑制劑)為間歇性跛行 Class I 建議。', False)],
    [('Pentoxifylline 已不再建議。', True)],
    [('多專科團隊(專科醫師、傷口照護、營養師)合作可降低罹病率與改善長期結果。', False)],
]
for it in items:
    add_bullet(doc, it)

add_heading(doc, '六、處置與後續', level=2)
items = [
    [('ALI 或慢性缺血急性惡化 → ', False), ('觀察、住院或轉院至具血管外科能力之中心', True), ('。', False)],
    [('慢性 PAD 無立即威脅 + 無其他急性疾病 → 可出院,門診追蹤血管外科或初級照護。', False)],
    [('衛教:症狀惡化(尤其疼痛)立即返診;若無禁忌,', False), ('aspirin 325 mg 起始,後續每日 81 mg', True), ('。', False)],
]
for it in items:
    add_bullet(doc, it)

add_heading(doc, '七、特殊族群:上肢缺血', level=2)
items = [
    [('上肢動脈急性阻塞 ', False), ('遠少於下肢', True), ('。', False)],
    [('因肩、肘周邊側支循環發達,缺血耐受性較佳;', False), ('靜止痛或壞疽罕見,除非有遠端栓塞', True), ('。', False)],
    [('原因:血管痙攣、動脈炎、外傷、粥狀硬化斑塊破裂、栓塞、醫源性損傷(如肱動脈導管)、胸廓出口症候群、動脈瘤、高凝狀態。', False)],
    [('檢查:臨床評估、肘上下節段血壓、Doppler、duplex US、動脈攝影。', False)],
    [('威脅肢體之手與前臂缺血治療:', False), ('Heparin + 緊急外科 thromboembolectomy', True), ('。', False)],
]
for it in items:
    add_bullet(doc, it)

doc.add_page_break()

# ============================================================
# Summary
# ============================================================
add_heading(doc, '★ 三章共通臨床要點(住院醫師考點)', level=1)
points = [
    [('主動脈剝離治療順序', True), (' :必先 ', False), ('β-blocker', True), (' ,再用血管擴張劑(否則反射性心跳加快增加 shear force)。', False)],
    [('AAA 不穩定病人 ', False), ('不可送出急診做 CT', True), (' ,先床邊 US + 立即會診。', False)],
    [('AAA 復甦採 ', False), ('允許性低血壓', True), (' ,SBP 目標 80–90 mmHg ,意識為指標。', False)],
    [('ALI 六大 P:缺一不可排除缺血', True), ('。', False)],
    [('Rutherford IIb', True), (' :', False), ('立即再血管化', True), (' ,不再延遲做影像。', False)],
    [('ALI 第一步:', False), ('Heparin bolus + infusion', True), ('。', False)],
    [('D-dimer 單獨不能排除主動脈剝離', True), (' (false negative 高達 18%)。', False)],
    [('CXR 不能排除 AAA 或剝離', True), ('(12–37% 完全正常)。', False)],
    [('AAA 易誤診為腎結石', True), (' ,需高度警覺。', False)],
    [('Aortoenteric fistula:消化道出血 + 主動脈移植史 → 必須懷疑。', True)],
]
for it in points:
    add_bullet(doc, it)

# Save
output_path = '/home/user/my-first-claude-project-/Tintinalli_Ch59-61_中文筆記.docx'
doc.save(output_path)
print(f'Saved: {output_path}')
